"""
Generates Pause-Health.ai-Build-Journal.docx -- a chronological record
of the prototype's construction, distilled from the full agent-chat
transcript that produced the codebase.

Output: <repo-root>/Pause-Health.ai-Build-Journal.docx

Run:
    .venv/bin/python scripts/build_journal.py
"""

from __future__ import annotations

import os
import sys
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).resolve().parent.parent
# The build journal lives outside this repo, alongside the Pause-Health.ai
# project deliverables. Override with PAUSE_JOURNAL_PATH if needed.
DEFAULT_OUTPUT_DIR = Path("/Users/maggie.hu/Projects/pause-health.ai")
OUTPUT_PATH = Path(
    os.environ.get(
        "PAUSE_JOURNAL_PATH",
        DEFAULT_OUTPUT_DIR / "Pause-Health.ai-Build-Journal.docx",
    )
)


# ---------- content ------------------------------------------------------

DOC_TITLE = "Pause-Health.ai — Build Journal"
DOC_SUBTITLE = (
    "A chronological record of how the prototype, investor brief, and "
    "multi-agent control plane were built — distilled from the full agent "
    "engineering chat."
)
DOC_AUTHOR_LINE = (
    "Authored from the Cursor agent transcript that produced the codebase. "
    "Compiled " + date.today().isoformat() + "."
)

OVERVIEW_PARAGRAPHS = [
    "Pause-Health.ai is an AI-led care platform for women navigating "
    "menopause. The repository in this journal contains everything built "
    "to date: the marketing site, the investor brief, the clickable "
    "prototype, the Python wearable-ingest worker, the MuleSoft mock "
    "integration plane, a Model Context Protocol server, a four-agent "
    "control plane with Google A2A handoff and a mocked MuleSoft Agent "
    "Fabric, and most recently a Salesforce Data 360 grounding layer "
    "feeding the Care Router.",
    "Each section of this journal corresponds to one focused build phase. "
    "Phases are listed in the order they happened. For each phase the "
    "journal records the ask, the decision points and trade-offs, what "
    "was actually built, and how it was verified. Where a phase changed "
    "the architecture or the investor narrative, that change is called "
    "out explicitly.",
    "Two ground rules across every phase. First, nothing in the prototype "
    "lies about what is real and what is mocked — every mocked surface "
    "carries a clear note in the API meta payload and a prototype-vs-"
    "production table on the corresponding investor page. Second, real "
    "credentials are never required for the prototype to demonstrate "
    "value; every external integration (Anthropic, Salesforce, MuleSoft, "
    "JupyterHealth) is gated behind an environment variable with a "
    "deterministic fallback so a reviewer can run the entire flow with "
    "no setup.",
]

ARCH_DIAGRAM = """\
                          Pause-Health.ai end-state architecture

   ┌───────────────────────────────────────────────────────────────────────┐
   │  Marketing site & investor brief        Clickable prototype           │
   │  (/, /about, /proposal/*)                (/demo/intake -> /demo/...)   │
   └───────────────────────────────────────────────────────────────────────┘
                                  │
                                  │  Patient submits intake
                                  ▼
   ┌───────────────────────────────────────────────────────────────────────┐
   │  Agentforce Service Agent (front door) — LIVE**                       │
   │  Pause_Health_Intake_Agent (Service Cloud) embedded via V2 Messaging  │
   │  for Web on pause-health.ai. Scripted Pause fallback for unconfigured │
   │  forks/previews.                                                      │
   └───────────────────────────────────────────────────────────────────────┘
                                  │
                                  │  Google A2A `tasks/send` with intake + grounding
                                  ▼
   ┌────────────┐                 ┌────────────────────────────────────┐
   │ Salesforce │ <───── REST ─── │ Pause Care Router (Anthropic Claude)│
   │ Health     │  IR + grounding │ Real Claude when ANTHROPIC_API_KEY  │
   │ Cloud      │ federated query │ set; deterministic policy fallback  │
   │ (LIVE)*    │                 │                                     │
   └────────────┘                 └────────────────────────────────────┘
       │ * Real OAuth 2.0 client credentials against a connected
       │   Salesforce Developer Edition org. SOQL against Contact +
       │   CareProgramEnrollee + CarePlan + Case. Falls back to the
       │   mocked Data 360 fixtures when SF_INSTANCE_URL / SF_CLIENT_ID
       │   / SF_CLIENT_SECRET are unset (zero-credential default).
                                  │
                                  │  Tool calls over MCP
                                  ▼
   ┌───────────────────────────────────────────────────────────────────────┐
   │  Pause MCP server (@pause-health/mcp)                                 │
   │  Four tools, each backed by a mocked MuleSoft Experience API         │
   └───────────────────────────────────────────────────────────────────────┘
                                  │
                                  │  REST over the integration plane
                                  ▼
   ┌───────────────────────────────────────────────────────────────────────┐
   │  MuleSoft Experience APIs (mocked)                                    │
   │   /api/mulesoft/health, /patient/[id]/timeline, /patient/[id]/intake, │
   │   /providers                                                          │
   └───────────────────────────────────────────────────────────────────────┘
                                  │
                                  │  Federated to / written to
                                  ▼
   ┌───────────────────────────────────────────────────────────────────────┐
   │  Clinical substrate: JupyterHealth Exchange (FHIR R5) + DBDP feature  │
   │  pipeline (FLIRT-based wearable HRV ingest)                           │
   └───────────────────────────────────────────────────────────────────────┘

   Cross-cutting:
     • MuleSoft Agent Fabric (mock) governs every agent — registry,
       policies, end-to-end traces, /demo/agent-fabric live console.
     • Salesforce Data 360 sits as a fifth agent on the Fabric and
       grounds the Care Router with longitudinal calculated insights.
       The Health Cloud layer (Phase 1) is LIVE against a real org;
       the Data Cloud unified-profile layer (Phase 2) is provisioned
       but deferred — 31 unified DMOs exist with no Data Streams.
     • A real Anypoint Platform org is connected for future MuleSoft
       work; today every Experience API is still served by the Next.js
       mock. Next-session plan lives in docs/MULESOFT_RUNBOOK.md.
     • ** Agentforce embedded chat (Phase 3) shipped 2026-06-02. Real
       Salesforce Agentforce Service Agent on Service Cloud, V2
       Messaging for Web bootstrap, routed through Omni-Channel
       (Agentforce Service Agent routing) directly to the agent.
       See docs/PHASE_3_RUNBOOK.md for the deployment topology and
       the two gotchas we hit (clientVersion=WebV1 default needing a
       Tooling API v65 PATCH; Messaging Channel routing needing to
       point at the agent, not a legacy Omni-Flow).
"""

MONOREPO_LAYOUT = """\
pause-health.ai/                            ← repo root (Apache-2.0, public on GitHub)
├── frontend/                               ← Next.js 14 App Router site
│   ├── app/                                  Pages + API routes
│   │   ├── (marketing routes, about, blog, careers, hipaa, etc.)
│   │   ├── proposal/*                        Investor brief deep-dives (16 routes)
│   │   ├── demo/*                            Clickable prototype (intake/routing/
│   │   │                                       patient/analytics/agent-fabric)
│   │   ├── provider/                         Browseable directory (Phase-2 NPPES);
│   │   │                                       /provider/[npi] profile pages
│   │   ├── changelog/, roadmap/              Versioned project state
│   │   └── api/
│   │       ├── agents/care-router/           A2A agent + tool calls
│   │       ├── agent-fabric/                 Console-backing endpoints
│   │       ├── intake/route-to-care-router/  Handoff orchestrator
│   │       ├── intake/prechat-context/       Hidden-prechat dossier endpoint
│   │       ├── mulesoft/                     Experience APIs (live-or-mock proxy)
│   │       ├── data-360/                     Health Cloud + Data Cloud read paths
│   │       ├── salesforce/headless-360/*     PKCE External Client App seam (6 routes)
│   │       ├── mcp/                          Streamable HTTP MCP server (Agentforce 3.0
│   │       │                                   Registry); /api/mcp/whoami diagnostic
│   │       └── agent-fabric/sf-sink/         Platform Event egress diagnostics
│   ├── components/                           UI + intake fallback + agentforce-embed +
│   │                                          recommended-providers + status-pill
│   ├── lib/                                  a2a, care-router, agent-fabric, data-360,
│   │   │                                      mulesoft-mocks, mscp-overlay, zip-centroids
│   │   ├── salesforce/                       auth.ts + grounding.ts + tests
│   │   ├── salesforce-headless360.ts         PKCE seam + validateMcpApiBearer
│   │   ├── salesforce-platform-event-sink.ts Agent Fabric trace egress
│   │   ├── mulesoft/health.ts + providers.ts Live-or-mock proxy with Auth0-JWT path
│   │   └── mcp/host.ts                       Care Router MCP-host shim
│   ├── scripts/                              smoke-test.mjs (168 routes),
│   │                                          salesforce-seed.mjs, salesforce-smoke.mjs,
│   │                                          grounding-smoke.mjs, preflight.mjs
│   └── public/.well-known/mcp.json           MCP discovery descriptor
│
├── cli/                                    ← @pause-health/cli (Headless 360 gap #4)
│   └── src/cli.ts + src/commands/{health,providers,timeline,intake}.ts
│
├── mcp/                                    ← @pause-health/mcp stdio MCP server
│   ├── src/server.ts                         Four MCP tools (stdio transport)
│   └── scripts/smoke.mjs                     End-to-end SDK smoke test
│
├── mulesoft/                               ← MuleSoft assets
│   ├── pause-mulesoft-health-v1/             DEPLOYABLE Mule 4 app on CloudHub 2.0
│   │                                          (v1.0.5, serving /health + /providers)
│   ├── pause-omh-to-fhir-library/            Phase-3 shared DataWeave library on Exchange
│   ├── specs/                                Eight OAS-3.0 spec assets on Exchange:
│   │                                          pause-jhe-system-api-spec,
│   │                                          pause-dbdp-system-api-spec,
│   │                                          pause-oura/whoop/garmin/healthkit/empatica-system-api-spec,
│   │                                          pause-ingest-process-api-spec
│   ├── flex-gateway/                         Docker + ngrok runtime enforcement
│   │                                          (JWT + rate limiting via Flex Gateway)
│   ├── flows/pause-process-api.example.xml   Reference Process-tier flow
│   └── transforms/omh-to-fhir.example.dwl    Reference DataWeave transform
│
├── salesforce/                             ← SFDX project (Phase 18b)
│   ├── force-app/main/default/{flows,objects,permissionsets,messagingChannels,
│   │   namedCredentials}                     Version-controlled Salesforce metadata
│   └── deploy.sh, retrieve.sh
│
├── pause_ingest/                           ← Python wearable ingest worker
│   ├── pause_ingest/{convert,exchange,features*,empatica,data_cloud,cohort}.py
│   ├── examples/oura_sample_upload.py        End-to-end smoke against real JHE
│   ├── examples/data_cloud_push.py           Push real DBDP features to Data Cloud
│   ├── tests/test_exchange_real_jhe.py       PAUSE_USE_REAL_JHE=1 opt-in suite
│   └── pyproject.toml
│
├── provider_ingest/                        ← NPPES → Pause directory pipeline
│   └── tracked refresh harness + 3-state sanctions overlay (CA/NY/TX)
│
├── jhe-local/                              ← JupyterHealth Exchange local stack
│   └── bootstrap.sh + teardown.sh            Docker-compose, idempotent seed
│
├── data-cloud/                             ← Salesforce Data Cloud assets
│   ├── Pause_HRV_RMSSD_30d.sql, Pause_Vasomotor_Burden_30d.sql,
│   │   Pause_Sleep_Disruption_7d.sql         Three live Calculated Insights
│   └── Pause_Wearable_Feature.dlo-schema.json
│
├── docs/                                   ← 18+ engineering runbooks (point-in-time)
│   ├── MULESOFT_RUNBOOK.md, MULESOFT_PHASE_1_HANDOFF.md, MULESOFT_API_MANAGER_RUNBOOK.md,
│   │   MULESOFT_PHASE_2_DATA_CLOUD.md, FLEX_GATEWAY_RUNBOOK.md
│   ├── HEADLESS_360_RUNBOOK.md, AGENTFORCE_VOICE_RUNBOOK.md,
│   │   AGENTFORCE_PROVIDER_ACTION_RUNBOOK.md, SF_PLATFORM_EVENT_SINK_RUNBOOK.md
│   ├── PHASE_2_ACTIVATION_CHECKLIST.md, PHASE_2_INGESTION_API_RUNBOOK.md,
│   │   PHASE_3_RUNBOOK.md
│   ├── JHE_SETUP_RUNBOOK.md, JHE_REAL_RUN_2026-06-16.md
│   ├── PROVIDER_GRAPH_PHASE_1_RUNBOOK.md, branch-protection.md
│   └── jupyterhealth-integration.md, mulesoft-integration.md
│
├── lighthouse-history/                     ← Nightly auto-maintained CWV time series
├── LICENSE, CONTRIBUTING.md, CODE_OF_CONDUCT.md, SECURITY.md  ← OSS hygiene (Apache-2.0)
└── README.md                                 ← Updated each phase
"""

PHASES = [
    {
        "title": "Phase 1 — Foundation: Next.js prototype, deploy, mobile nav",
        "ask": (
            "Bootstrap a credible marketing-quality Next.js site and an "
            "investor-ready prototype on top of the existing FastAPI repo. "
            "Edit the About page (HQ -> Irvine, CA; add Maggie C. Hu as "
            "Founder & CEO). Check the project into GitHub. Deploy to "
            "Vercel. Make the top navigation mobile-friendly."
        ),
        "decisions": [
            "Kept the FastAPI app and the new Next.js frontend side-by-side "
            "in one monorepo rather than splitting repos. The legacy app is "
            "still useful for tests; the new product surface lives in "
            "frontend/.",
            "Vercel application preset switched from FastAPI to Next.js so "
            "the right framework detection runs. The 404 after first deploy "
            "was the result of the wrong preset; once switched, the App "
            "Router routes resolved.",
            "Mobile nav implemented as a separate client component "
            "(MobileNav) so the desktop header remains a server component. "
            "Hamburger toggles a portal-style drawer instead of stacking "
            "links on top of content.",
        ],
        "built": [
            "frontend/ — Next.js 14 App Router, marketing pages, footer, "
            "manifest, sitemap, robots, OpenGraph, Twitter card metadata, "
            "SEO basics.",
            "GitHub repository at github.com/hucmaggie/pause-health.ai.",
            "Vercel deployment with the Next.js preset and a working "
            "production URL.",
            "components/mobile-nav.tsx with a responsive hamburger drawer.",
        ],
        "verified": [
            "ESLint clean across the frontend.",
            "Lighthouse nightly workflow scheduled via GitHub Actions.",
            "Manual smoke check of every top-level page on both desktop "
            "and mobile widths.",
        ],
    },
    {
        "title": "Phase 2 — Investor brief: /proposal/full and deep-dives",
        "ask": (
            "Expand the prototype with detailed investor-focused content. "
            "Cover customer selection (health system + payer profiles), "
            "key customer insights from research, data inventory and "
            "strategy, competition, digital strategy, and technology "
            "choices. Convert the single long page into a /proposal/full "
            "route that shares the layout, with clean prose styling."
        ),
        "decisions": [
            "Composed every investor page through a shared ProposalShell "
            "component so the side navigation, eyebrow/title/subtitle "
            "rhythm, and prose styling stay consistent.",
            "Standardized on three primitives across every investor page: "
            "card-grid (for items), table-wrap (for prototype-vs-production "
            "comparisons), and metric-list (for takeaways). Avoids ad-hoc "
            "layouts.",
            "The /proposal page itself became an index with one-paragraph "
            "summaries pointing into each deep-dive, so reviewers can "
            "skim or drill in without backtracking.",
        ],
        "built": [
            "/proposal index page with summarized links to every deep-dive.",
            "/proposal/full — the single-document narrative for a complete "
            "investor read.",
            "/proposal/customers, /proposal/insights, /proposal/data, "
            "/proposal/competition, /proposal/strategy, /proposal/technology.",
        ],
        "verified": [
            "ESLint clean.",
            "next build passes; sitemap regenerated to include all new "
            "investor routes.",
        ],
    },
    {
        "title": "Phase 3 — JupyterHealth integration design (Option C)",
        "ask": (
            "Figure out how to integrate JupyterHealth via their GitHub "
            "organization. Choose the deepest viable integration option."
        ),
        "decisions": [
            "Option C selected: compose with three JupyterHealth projects "
            "rather than picking one — JupyterHealth Exchange (the FHIR R5 "
            "+ Open mHealth data plane), omh-shim (wearable -> OMH "
            "conversion), and jupyterhealth-client (Python API client).",
            "Treated JupyterHealth as Pause's clinical substrate rather "
            "than as 'one of N integrations' — every wearable signal lands "
            "in FHIR via JupyterHealth, and the Care Router's longitudinal "
            "view federates over the same store.",
            "Authored an investor page that frames JupyterHealth as the "
            "open-source data substrate so that the customer's data team "
            "sees Pause as additive, not as a competing data platform.",
        ],
        "built": [
            "/proposal/integration investor page covering the three-piece "
            "composition, the FHIR R5 substrate, and the data plane.",
            "Design doc identifying the contract points: how OMH payloads "
            "from wearables become FHIR Observations posted to the "
            "JupyterHealth Exchange, and how the Python client reads them "
            "back for feature computation.",
            "README updated to enumerate the JupyterHealth pieces.",
        ],
        "verified": [
            "Investor page renders, links to upstream JupyterHealth repos, "
            "and is referenced from the /proposal index.",
        ],
    },
    {
        "title": "Phase 4 — DBDP wearable ingest (Steps 1 and 2)",
        "ask": (
            "Integrate with DBDP's code repository for wearable ingest — "
            "research the ecosystem, identify FLIRT and a particular HRV "
            "script, and integrate them into a Python ingest worker. Use "
            "both synthetic and DHDR-style test data."
        ),
        "decisions": [
            "Step 1 scaffolded a pause_ingest/ Python package on top of "
            "FLIRT (DBDP's feature-generation toolkit). FLIRT handles "
            "windowing, sleep detection, and HRV feature extraction; "
            "Pause wraps it for the menopause-specific feature set.",
            "Step 2 ported the DBDP HRV calculator and an Empatica E4 "
            "ingestion path. Devicely was deferred — it does not yet "
            "support Python 3.13, and forcing 3.11 would have churned the "
            "rest of the toolchain.",
            "Tests run against both synthetic fixtures and a DHDR-style "
            "sample so reviewers can verify the pipeline with no external "
            "data set.",
        ],
        "built": [
            "pause_ingest/ Python package — FLIRT-based feature pipeline, "
            "HRV calculator, Empatica E4 reader, synthetic + DHDR-style "
            "test fixtures, pytest suite, pyproject.toml.",
            "/proposal/dbdp investor page covering the DBDP ecosystem, "
            "FLIRT integration, the calculated features Pause needs, and "
            "the deferred devicely work.",
            "README updated to describe the wearable ingest worker.",
        ],
        "verified": [
            "Python tests green (synthetic + DHDR fixtures).",
            "Investor page renders and cross-links from /proposal/integration.",
        ],
    },
    {
        "title": "Phase 5 — The Menopause Society and provider graph",
        "ask": (
            "Integrate with the National Menopause Society's data and "
            "provider list. Implement Path C (deep-link referral) "
            "immediately and document Path B (NPPES-derived provider "
            "graph) for investors."
        ),
        "decisions": [
            "Path C ships now: a deterministic deep-link to the "
            "Menopause Society's MSCP-credentialed clinician directory "
            "with the patient's filters preserved. Zero data-licensing "
            "risk, immediate clinical value, real signal in the demo.",
            "Path B documented: a Pause-owned provider graph derived "
            "from CMS NPPES (NPI Registry) data, joined to MSCP "
            "credential status. This becomes the long-term competitive "
            "moat; the investor page lays out the build plan rather than "
            "shipping a half-built directory now.",
            "Provider mock data reused in the MuleSoft providers Experience "
            "API so the prototype already exercises the shape Pause will "
            "ship in production.",
        ],
        "built": [
            "/proposal/menopause-society — Path C live, Path B documented.",
            "/proposal/provider-graph — the NPPES-derived graph plan, the "
            "competitive moat argument, the phased build.",
            "Mock provider directory used by /api/mulesoft/providers.",
        ],
        "verified": [
            "Deep-link Path C resolves to the Menopause Society directory "
            "with filters preserved.",
            "Mock provider API returns deterministic results filterable "
            "by ZIP and menopause-experience flag.",
        ],
    },
    {
        "title": "Phase 6 — Agentforce Service Agent (Option A)",
        "ask": (
            "Integrate Salesforce Agentforce Service Agent for patient "
            "intake. Implement Option A: a real Embedded Messaging "
            "Service Agent with the Pause-branded palette, plus a "
            "graceful scripted fallback if credentials are not "
            "configured."
        ),
        "decisions": [
            "Real Salesforce Embedded Messaging for Web (Enhanced Chat v2 "
            "inline mode) loaded when the Vercel environment variables are "
            "set; scripted React fallback runs otherwise. One environment "
            "flag, two code paths, identical UX shape.",
            "All Salesforce configuration kept as Vercel environment "
            "variables — no secrets in the repo. Documented in .env.example.",
            "The intake completion event is the trigger for everything "
            "downstream (Care Router handoff, Data 360 grounding). The "
            "fallback emits the same event so the rest of the multi-agent "
            "trace works identically whether Salesforce is configured or not.",
        ],
        "built": [
            "components/agentforce-fallback.tsx — scripted intake mirroring "
            "the real agent's question set.",
            "Real Agentforce embed gated by Vercel env vars.",
            "/proposal/agentforce investor page covering Embedded "
            "Messaging, configuration model, fallback strategy.",
            ".env.example updated.",
        ],
        "verified": [
            "Fallback runs end-to-end with no env vars set.",
            "Real Agentforce embed renders correctly when env vars are set "
            "(verified against a Salesforce trial org).",
        ],
    },
    {
        "title": "Phase 7 — MuleSoft integration plane (Option C)",
        "ask": (
            "Add MuleSoft to the prototype for integrations with "
            "JupyterHealth and DBDP. Option C: design doc + investor "
            "page + reference Mule/DataWeave files + mocked API "
            "endpoints, with cross-links from existing investor pages."
        ),
        "decisions": [
            "API-led connectivity adopted as the contract: System APIs "
            "wrap raw sources (JupyterHealth FHIR, DBDP feature store, "
            "EHR-of-record); Process APIs orchestrate validation and "
            "writes; Experience APIs serve the agents and the UI.",
            "DataWeave omh-to-fhir.dwl included as a real artifact, not a "
            "screenshot — reviewers can read the transform end to end.",
            "Mocked Experience APIs centralized in frontend/lib/"
            "mulesoft-mocks.ts so the same fixtures back the Experience "
            "API routes and the MCP server in the next phase.",
        ],
        "built": [
            "mulesoft/ directory — system-api, process-api, experience-api "
            "reference scaffolds, plus dataweave/omh-to-fhir.dwl.",
            "frontend/lib/mulesoft-mocks.ts — shared fixtures.",
            "Mocked Experience APIs: /api/mulesoft/health (kept for backward "
            "compatibility), /api/mulesoft/patient/[id]/timeline, "
            "/api/mulesoft/patient/[id]/intake, /api/mulesoft/providers.",
            "/proposal/mulesoft investor page covering API-led tiers, "
            "DataWeave, deployment topology, and the prototype-vs-production "
            "table.",
            "Cross-links added from /proposal/integration and /proposal/dbdp.",
        ],
        "verified": [
            "curl smoke tests against all four mocked Experience APIs.",
            "ESLint clean; next build passes with the new routes.",
        ],
    },
    {
        "title": "Phase 8 — Pause MCP server",
        "ask": (
            "Add MuleSoft MCP servers for the mock APIs. Resolved with "
            "the user: this is really 'expose the MuleSoft Experience "
            "APIs as MCP tools so Claude Desktop, Cursor, and Agentforce "
            "can call them.' Build a real MCP server, three new mock "
            "Experience API routes, a .well-known descriptor, and an "
            "investor page (Option C — the most complete option)."
        ),
        "decisions": [
            "Standalone Node package mcp/ rather than embedding the MCP "
            "server in Next.js. The MCP transport is stdio, not HTTP — "
            "embedding it would have been the wrong shape.",
            "Four tools, one per Experience API: get_patient_timeline, "
            "get_patient_intake, find_menopause_providers, "
            "experience_api_health. Inputs validated with Zod.",
            "Published /.well-known/mcp.json from the Next.js public "
            "directory so any MCP client can auto-discover Pause's tools.",
            "Smoke test (mcp/scripts/smoke.mjs) drives the server through "
            "the official MCP SDK client to prove the tools resolve "
            "end-to-end.",
        ],
        "built": [
            "mcp/ package — package.json, src/server.ts, "
            "scripts/smoke.mjs, README.md, tsconfig.json.",
            "Three new mocked Experience APIs in frontend/app/api/mulesoft/ "
            "to back the new MCP tools.",
            "frontend/public/.well-known/mcp.json discovery descriptor.",
            "/proposal/mcp investor page covering the four tools, why MCP, "
            "client registration snippets (Claude Desktop, Cursor), "
            "prototype-vs-production table, phased plan.",
        ],
        "verified": [
            "npm run smoke against the built MCP server: list-tools and "
            "every individual tool call return expected payloads.",
            "Discovery descriptor served at /.well-known/mcp.json with the "
            "right content type.",
            "MCP client registration tested with Cursor's mcp.json format.",
        ],
    },
    {
        "title": "Phase 9 — Multi-agent control plane: Anthropic + A2A + Agent Fabric (L3)",
        "ask": (
            "Add an Anthropic agent that processes care details and "
            "routes the case to the appropriate pathway, showing the A2A "
            "protocol. Add MuleSoft Agent Fabric to showcase multi-agent "
            "orchestration, monitoring, security, and governance. L3 "
            "implementation: investor page, mocked Agent Fabric console, "
            "and a real end-to-end A2A handoff from intake to Care Router "
            "with traces appearing live in the console."
        ),
        "decisions": [
            "Anthropic Care Router gated by ANTHROPIC_API_KEY: real Claude "
            "Sonnet 4.5 when the key is set, deterministic policy engine "
            "as a graceful fallback. Same decision shape from either path.",
            "Google A2A used for agent-to-agent handoff (Agentforce -> "
            "Care Router). MCP retained as the tool-call surface to the "
            "data plane. Two open protocols, two different concerns — "
            "documented as such on the investor page.",
            "MuleSoft Agent Fabric implemented as an in-memory mock: agent "
            "registry, policy catalog, ring-buffer trace store, governance "
            "evaluator. Module-scoped global so traces persist across "
            "Next.js dev hot reloads and every API route in the same Node "
            "process sees the same store.",
            "/demo/agent-fabric became a live console: polls the registry, "
            "shows policy catalog, lists recent tasks, renders parent/child "
            "spans for any taskId. Test-case buttons drive real A2A flows.",
            "Twelve governance policies authored covering PHI, model "
            "allow-list, clinical rationale requirement, MCP allow-list, "
            "audit, network, FHIR substrate — each tagged enforced/advisory "
            "and block/audit.",
        ],
        "built": [
            "frontend/lib/a2a.ts — A2A protocol types and helpers.",
            "frontend/lib/care-router.ts — scripted policy engine plus "
            "Anthropic Claude integration.",
            "frontend/lib/agent-fabric.ts — registry, policies, trace "
            "store, governance evaluator.",
            "/api/agents/care-router/.well-known/agent.json — A2A Agent "
            "Card discovery document.",
            "/api/agents/care-router/tasks — A2A tasks/send endpoint.",
            "/api/agent-fabric/agents, /policies, /traces, "
            "/governance/evaluate — console-backing endpoints.",
            "/api/intake/route-to-care-router — handoff orchestrator that "
            "stitches Agentforce intake to the Care Router over A2A.",
            "components/latest-care-router-decision.tsx — polls traces and "
            "shows the most recent decision on /demo/routing.",
            "/demo/agent-fabric — the live multi-agent console.",
            "/proposal/agent-fabric investor page covering the four agents, "
            "the two protocols, the Fabric's role, prototype-vs-production, "
            "and phased plan.",
            "Anthropic SDK added as a soft dependency (dynamic import).",
        ],
        "verified": [
            "End-to-end POST to /api/intake/route-to-care-router produces a "
            "RoutingDecision and a trace span tree visible at /demo/agent-"
            "fabric.",
            "Governance policies block red-flag-missing intakes; the trace "
            "shows the blocked span and the A2A task in failed state.",
            "Anthropic path verified with a live API key (still falls back "
            "gracefully when the key is removed).",
            "ESLint clean; next build passes.",
        ],
    },
    {
        "title": "Phase 10 — Code Repository link in the top nav",
        "ask": (
            "Add a 'Code Repository' link to the top navigation menu with "
            "the GitHub repository URL."
        ),
        "decisions": [
            "Added to both desktop (frontend/app/layout.tsx) and mobile "
            "(components/mobile-nav.tsx) navigation, behind external-link "
            "rendering with target=_blank and rel=noopener noreferrer.",
            "MobileNav's NavLink type extended with an `external?: boolean` "
            "flag so external links are handled uniformly going forward "
            "rather than as a special case for this one entry.",
        ],
        "built": [
            "frontend/app/layout.tsx — desktop nav.",
            "frontend/components/mobile-nav.tsx — external NavLink "
            "rendering.",
        ],
        "verified": [
            "Visual check via curl against the dev server on both desktop "
            "and mobile widths.",
            "ESLint clean.",
        ],
    },
    {
        "title": "Phase 11 — Salesforce Data 360 grounding (L2)",
        "ask": (
            "How could we incorporate Data 360 into the prototype? "
            "Decided on L2: investor page + mocked Data 360 read APIs + "
            "Care Router actually fetches grounding from Data 360 before "
            "deciding, with the fetch recorded as a trace span. Grounding "
            "fetched before the Care Router (cleanest trace). Leave "
            "/demo/patient alone for now but add a 'View Data 360 record' "
            "link from the Agent Fabric console."
        ),
        "decisions": [
            "Data 360 is the unified-read plane on top of the existing "
            "integration plane (MuleSoft) and the clinical substrate "
            "(JupyterHealth + DBDP). Composes with everything else; "
            "replaces nothing.",
            "Grounding fetched in /api/intake/route-to-care-router before "
            "the A2A handoff, attached to the A2A message as a data part. "
            "The Care Router does not need to know about Data 360 — it "
            "just consumes more context.",
            "Care Router updated (both scripted and Claude paths) to read "
            "the grounding and cite specific calculated insights in its "
            "rationale. A virtual MSCP visit gets promoted to in-person "
            "when vasomotor burden >= 60 or cohort percentile >= 75 — a "
            "visible, defensible decision change driven by the grounding.",
            "Data 360 registered as the fifth agent on the Agent Fabric, "
            "with three new block-enforced policies: zero-copy federation, "
            "consent required before grounding, segment activation "
            "allow-list. Audit policy extended to cover Data 360.",
            "Federated-record JSON deep-link added to /demo/agent-fabric "
            "so reviewers can see the underlying patient record for any "
            "traced task in one click.",
        ],
        "built": [
            "frontend/lib/data-360.ts — federated patient store, four "
            "calculated insights (HRV z-score, vasomotor burden, sleep, "
            "days-since-MSCP), longitudinal observations, cohort "
            "comparison with pathway-outcome breakdown, identity "
            "resolution stub, four population segments.",
            "/api/data-360/patient/[id]/grounding — the bundle the Care "
            "Router consumes.",
            "/api/data-360/patient/[id]/record — full federated record.",
            "/api/data-360/segments — population segment catalog.",
            "/api/data-360/identity/resolve — IR endpoint.",
            "Updated /api/intake/route-to-care-router to fetch IR + "
            "grounding before the A2A handoff and record two new spans.",
            "Updated frontend/lib/care-router.ts to consume grounding "
            "and surface it on every decision.",
            "Updated /api/agents/care-router/tasks to pass grounding "
            "through to route().",
            "/proposal/data-360 investor page — zero-copy federation, "
            "four-span trace, prototype-vs-production table, four-phase "
            "plan, investor takeaways.",
            "Wired into ProposalShell nav, /proposal index, sitemap; "
            "cross-linked from /proposal/agent-fabric, /proposal/mulesoft, "
            "/proposal/agentforce.",
            "README updated.",
        ],
        "verified": [
            "ESLint clean across 17 touched files.",
            "next build passes with 54 routes (4 new Data 360 endpoints + "
            "the new investor page).",
            "End-to-end smoke test: POST a moderate hot-flash intake -> "
            "decision routes to mscp-in-person (correctly promoted from "
            "virtual by grounding) -> trace shows four spans "
            "(intake.complete -> data360.identity.resolve -> "
            "data360.grounding.federated-query -> a2a.tasks/send) with "
            "parent/child correlation intact and grounding insights cited.",
            "Agent Fabric registry shows Data 360 with data-grounding "
            "governance tier; three Data 360 policies showing as "
            "block/enforced.",
        ],
    },
    {
        "title": (
            "Phase 12 — Real Salesforce org integration: Health Cloud "
            "grounding (Full Path A, Phase 1)"
        ),
        "ask": (
            "User stated: 'I have a Salesforce org that I could hook up "
            "the prototype with.' Phase 11 Data 360 was entirely mocked; "
            "this pivot moved the grounding plane from in-memory "
            "fixtures to a real, connected Salesforce Developer Edition "
            "org. Scope agreed: Full Path A (Health Cloud objects + Data "
            "Cloud unified profile + real Agentforce intake), executed "
            "as three sub-phases. Phase 12 covers sub-phase 1: real "
            "Health Cloud grounding."
        ),
        "decisions": [
            "Authentication: OAuth 2.0 Client Credentials Flow via "
            "Salesforce External Client App (modern replacement for the "
            "legacy Connected App). Server-to-server only; no user "
            "redirect, no browser dependency. The same pattern carries "
            "over to MuleSoft, Anypoint, and any future server-to-server "
            "integration.",
            "Real data: six menopause-specific Health Cloud records "
            "seeded into the org via a dedicated idempotent script "
            "(scripts/salesforce-seed.mjs). Used Contact + CarePlan + "
            "CareProgramEnrollee + Case rather than fabricating custom "
            "objects, so the demo composes with any future Health Cloud "
            "package without remapping.",
            "Schema reality forced two adjustments: linked "
            "CareProgramEnrollee to Account (not Contact); dropped the "
            "non-existent EnrolleeType column. The seeding script "
            "introspects the org's actual schema rather than assuming "
            "any documented field.",
            "Graceful degradation preserved: when SF_INSTANCE_URL / "
            "SF_CLIENT_ID / SF_CLIENT_SECRET are unset, the API routes "
            "fall back to the deterministic Data 360 mock. Reviewers, "
            "Vercel previews, and CI run with zero credentials. When set, "
            "every Care Router decision is grounded in real org data and "
            "the Agent Fabric trace span carries _source: 'real'.",
            "Zscaler hit once on the way: the corporate proxy intercepted "
            "*.c360a.salesforce.com (Data Cloud's hostname family) with a "
            "504 Gateway Timeout. Pausing Zscaler resolved it; documented "
            "as a known constraint for any future Salesforce-edge work.",
        ],
        "built": [
            "frontend/lib/salesforce/auth.ts — OAuth client credentials "
            "token acquisition, in-memory caching with expiry, request "
            "deduplication, isSalesforceConfigured() guard.",
            "frontend/lib/salesforce/grounding.ts — real SOQL fetcher "
            "and identity resolver against Contact + CarePlan + "
            "CareProgramEnrollee + Case; preserves the GroundingContext "
            "shape the Care Router already consumes.",
            "frontend/scripts/salesforce-smoke.mjs — standalone "
            "auth-and-query smoke test runnable independently of Next.js.",
            "frontend/scripts/salesforce-seed.mjs — idempotent demo-data "
            "seeder (six Contacts + linked CarePlan/Enrollee/Case rows) "
            "with explicit cleanup logic.",
            "frontend/scripts/grounding-smoke.mjs — end-to-end smoke of "
            "the real grounding fetcher.",
            "Wiring in /api/data-360/identity/resolve and "
            "/api/intake/route-to-care-router to use real Salesforce data "
            "when configured, mock otherwise.",
            "Agent Fabric console (/demo/agent-fabric) extended with a "
            "trace-level _source banner: every span shows 'real' or "
            "'mock' so the demo never lies about provenance.",
            "/proposal/data-360 investor page updated to mark Phase 1 "
            "as LIVE with a banner.",
            "README.md updated with the new live integrations section "
            "and the External Client App setup walkthrough.",
            ".env.example documented SF_INSTANCE_URL, SF_CLIENT_ID, "
            "SF_CLIENT_SECRET, SF_API_VERSION with comments explaining "
            "the OAuth client credentials policy and the My Domain URL "
            "requirement.",
        ],
        "verified": [
            "salesforce-smoke.mjs returns a valid access token against "
            "the connected org.",
            "salesforce-seed.mjs seeds six demo Contacts + linked "
            "Health Cloud records on a fresh run, then re-running "
            "produces zero diffs (idempotent).",
            "grounding-smoke.mjs returns a populated GroundingContext.",
            "End-to-end: POST a moderate hot-flash intake -> trace span "
            "for data360.grounding.federated-query shows _source: 'real' "
            "with citations pulled from the live org -> Care Router "
            "decision visibly changes vs the mocked-only path.",
            "When env vars are removed, the same flow re-runs and the "
            "trace span flips back to _source: 'mock' without code "
            "changes — graceful degradation verified both directions.",
        ],
    },
    {
        "title": (
            "Phase 13 — Polish: schema cleanup, auth unit tests, "
            "warn-once deduplication"
        ),
        "ask": (
            "Before moving to the next sub-phase of the Salesforce "
            "integration, polish three things the Phase 12 implementation "
            "left rough: (1) the seeded Contact.LastName was awkwardly "
            "encoded ('Pause Demo Patient:<name>'); (2) lib/salesforce/"
            "auth.ts had no unit tests; (3) the fallback path warned on "
            "every request when Salesforce was misconfigured, drowning "
            "the dev console."
        ),
        "decisions": [
            "Contact schema fix: chose the delete_reseed approach (user "
            "preference) rather than in-place update. New schema uses "
            "Contact.Title = 'Pause Demo Patient' and Contact.Department "
            "= 'Pause Demo' for demo tagging, with LastName carrying the "
            "actual surname. The cleanup predicate ORs across old and "
            "new tagging so a re-run safely deletes either generation.",
            "Auth tests with vitest covering happy-path, cached token "
            "reuse, in-flight deduplication, error propagation, expiry "
            "handling. Seventeen tests total.",
            "Warn-once helper (warnSalesforceDegradationOnce) keys on "
            "context + error name + error message prefix. Dev console "
            "stays usable even when the org is intentionally unconfigured "
            "for an entire session; the first failure of each category "
            "still surfaces.",
        ],
        "built": [
            "frontend/lib/salesforce/auth.test.ts — 17 unit tests "
            "covering getSalesforceConfig, isSalesforceConfigured, "
            "getAccessToken (cache hit, in-flight dedup, error paths, "
            "expiry).",
            "frontend/lib/salesforce/grounding.test.ts — 6 tests for "
            "warnSalesforceDegradationOnce dedup behavior.",
            "Updated scripts/salesforce-seed.mjs with the new tagging "
            "schema and a unified cleanup predicate.",
            "Updated scripts/grounding-smoke.mjs SOQL to match the new "
            "tagging.",
            "Updated /api/data-360/identity/resolve/route.ts and "
            "/api/intake/route-to-care-router/route.ts to use the "
            "warn-once helper.",
        ],
        "verified": [
            "npm test (vitest) passes including all 23 new salesforce/* "
            "tests.",
            "salesforce-seed.mjs delete + reseed cycle leaves zero "
            "stragglers in the org.",
            "Dev console during a Salesforce-unconfigured run shows one "
            "warn per failure category instead of one warn per request.",
            "Committed and pushed to main.",
        ],
    },
    {
        "title": (
            "Phase 14 — Salesforce Data Cloud (Phase 2): investigation "
            "and deferral"
        ),
        "ask": (
            "Move to sub-phase 2 of the Full Path A integration: real "
            "Data Cloud unified-profile layer feeding the Care Router."
        ),
        "decisions": [
            "Investigation confirmed the org's Data Cloud workspace is "
            "PROVISIONED but EMPTY: 31 unified DMOs exist "
            "(ssot__Individual__dlm etc.) with no Data Streams currently "
            "feeding them. Activation requires creating Data Streams, "
            "DLO mappings, Identity Resolution rules, and Calculated "
            "Insights — a 2-4 hour Setup UI exercise that produces no "
            "engineering artifact in the repo.",
            "Two OAuth scope additions were required to make Data Cloud "
            "even introspectable from outside Salesforce: cdp_api, "
            "cdp_profile_api, cdp_ingest_api on the External Client App. "
            "The /services/a360/token endpoint rejected calls without "
            "them.",
            "Decision: defer Phase 2 to a dedicated session with screen-"
            "share clickthrough. The engineering side is ready (lib/"
            "salesforce/auth.ts handles CDP scopes; the grounding fetcher "
            "is structured to accept a unified Individual record); the "
            "remaining work is entirely UI configuration inside the "
            "Salesforce org.",
        ],
        "built": [
            "Updated External Client App scopes documented in .env.example.",
            "README.md updated to mark Phase 2 as 'Data Cloud workspace "
            "provisioned but no Data Streams currently feed it; "
            "activation is a 2-4 hour Setup UI exercise.'",
            "/proposal/data-360 still accurately reflects Phase 1 LIVE / "
            "Phase 2 documented-but-deferred.",
        ],
        "verified": [
            "OAuth token request against /services/a360/token returns 200 "
            "with the new cdp_* scopes (was 'invalid_scope' before).",
            "SOQL against ssot__Individual__dlm and similar unified DMOs "
            "succeeds but returns zero rows — proving the schema exists "
            "and no Data Streams have hydrated it.",
            "No regression on Phase 1: Health Cloud grounding still "
            "shows _source: 'real' on every span.",
        ],
    },
    {
        "title": (
            "Phase 15 — Real Agentforce embedded chat (Phase 3): "
            "investigation, partial wiring, deferral with runbook"
        ),
        "ask": (
            "Pivot from Phase 2 to sub-phase 3: replace the scripted "
            "/demo/intake fallback with the real Salesforce Agentforce "
            "Embedded Messaging widget, served from a configured "
            "Experience Cloud site."
        ),
        "decisions": [
            "Investigation found the org already had: a Messaging "
            "channel, an Experience site, three agents, and an "
            "Embedded Service deployment named SDO_Messaging_for_Web. "
            "Plumbing this in took ~80% of the session.",
            "All four NEXT_PUBLIC_AGENTFORCE_* env vars set; "
            "components/agentforce-embed.tsx switched from inline to "
            "floating displayMode after extensive DOM debugging; "
            "globals.css updated with a compact launcher-callout and a "
            "minimal #embedded-messaging z-index defence.",
            "Hit a HARD CEILING that we don't control: the SDO sample "
            "deployment's runtime config endpoint (salesforce-scrt.com) "
            "blocks external origins via missing Access-Control-Allow-"
            "Origin, AND the Experience site's commcsp policy hardcodes "
            "frame-ancestors to a DIFFERENT org's domain. Neither is "
            "modifiable from the consumer side. The launcher mounts in "
            "the DOM but its inner iframe stays display:none after the "
            "config fetch is CORS-blocked.",
            "Three Salesforce-side fixes attempted and verified: added "
            "http://localhost:3000 to CorsWhitelistEntry; set "
            "AreGuestUsersAllowed = true via Tooling API (the field is "
            "not directly writable, had to PATCH the Metadata field); "
            "switched the Experience site deployment from V1 to V2. None "
            "lifted the SDO-sample restrictions.",
            "Decision: defer to a dedicated session where the user "
            "authors a Pause-Health-owned Embedded Service deployment "
            "+ Experience site + agent. Captured the full sequence in "
            "docs/PHASE_3_RUNBOOK.md so the next session is "
            "click-through-only rather than re-investigation.",
        ],
        "built": [
            "components/agentforce-embed.tsx — production-shape React "
            "wrapper around the V2 embedded_service_bootstrap, with "
            "floating displayMode and a status-driven launcher callout "
            "(loading / ready / error states).",
            "components/agentforce-fallback.tsx — refined Pause-branded "
            "scripted intake; emits the same intake.completed event as "
            "the real widget so downstream A2A + Care Router + Agent "
            "Fabric flows are unaffected.",
            "globals.css — agentforce-launcher-callout styles plus a "
            "minimal z-index override for #embedded-messaging so the "
            "Salesforce-injected launcher sits above Pause's toast "
            "region (no other overrides — Salesforce owns its own "
            "internal CSS).",
            "lib/agentforce.ts — getAgentforceConfig() helper that "
            "reads and validates the four NEXT_PUBLIC_AGENTFORCE_* env "
            "vars; returns null when any is missing (drives the "
            "fallback).",
            "docs/PHASE_3_RUNBOOK.md — the 2-4 hour next-session "
            "playbook with the exact Salesforce restrictions documented "
            "verbatim from the DevTools Console (CORS missing-header, "
            "commcsp frame-ancestors directive).",
            ".env.example and .env.local annotated with multi-paragraph "
            "comments explaining the SDO-sample restrictions, so the "
            "next agent or developer does not waste a session on the "
            "same dead end.",
            "README.md updated with the /demo/intake status section.",
        ],
        "verified": [
            "Bootstrap script downloads and embedded_service_bootstrap "
            ".init succeeds — the integration code is correct.",
            "DevTools Console reproduces the two specific Salesforce-"
            "side blockers (CORS and CSP frame-ancestors) deterministically.",
            "With env vars unset, /demo/intake renders the polished "
            "Pause-branded scripted fallback end-to-end.",
            "Committed and pushed: floating-mode wiring + investigation "
            "findings + next-session runbook.",
        ],
    },
    {
        "title": (
            "Phase 16 — MuleSoft Anypoint integration: investigation "
            "and next-session runbook"
        ),
        "ask": (
            "User stated: 'I have a MuleSoft org to be used for the "
            "prototype.' Scope agreed: investigate-only (~30 minutes), "
            "no Anypoint UI clickthrough, no commits to wiring. Produce "
            "a runbook calibrated against the Salesforce Phase 1 / "
            "Phase 3 lessons."
        ),
        "decisions": [
            "Inventoried the existing MuleSoft surface area: 4 mocked "
            "Experience-API routes, real Mule 4 reference flow, real "
            "DataWeave 2.0 transform (OMH -> FHIR R5), MCP server "
            "already supporting PAUSE_MCP_BASE_URL for swapping to a "
            "real Anypoint runtime, polished investor page, and full "
            "design doc. The story is more developed than Salesforce "
            "was at the same point — only the live runtime is missing.",
            "Probed Anypoint Platform hostnames. anypoint.mulesoft.com, "
            "eu1, gov, the OAuth token endpoint, us-e1.cloudhub.io, "
            "us-east-1.cloudhub.io, mq-us-east-1.anypoint.mulesoft.com, "
            "and the Exchange asset CDN all resolve and return their "
            "expected status codes. The Anypoint OAuth endpoint cleanly "
            "returns 401 to bogus client credentials — same shape as "
            "Salesforce, which means the auth pattern carries over.",
            "Headline network finding: Anypoint resolves through "
            "*.edge2.salesforce.com — the same edge family as the user's "
            "Salesforce org. Good news (Zscaler exception likely "
            "covers both), risk (if Zscaler tightens *.salesforce.com "
            "later, both Salesforce and MuleSoft break).",
            "Recommended scope: replace ONE mocked Experience API "
            "(/api/mulesoft/health) with a real Mule app deployed to "
            "CloudHub 2.0. Reasons: canonical demo URL linked from "
            "everywhere; payload fully synthetic so safe on a public "
            "worker; no upstream dependencies (the bundle does not "
            "need live JHE or pause_ingest); MCP server already supports "
            "the swap; tells the right investor story (one URL flips "
            "from 'mocked' to 'live on Anypoint Platform').",
            "Calibration against Salesforce: this looks like Phase 1 "
            "(server-to-server OAuth, no embedding, graceful "
            "degradation pattern proven) rather than Phase 3 (CORS/CSP-"
            "bound widget). High confidence the next session ships in "
            "one sitting.",
        ],
        "built": [
            "docs/MULESOFT_RUNBOOK.md (357 lines) — surface-area "
            "inventory, Zscaler probe results table, recommended scope, "
            "7-step playbook (Connected App -> Code Builder/Studio -> "
            "deploy -> wire with graceful degradation -> investor-page "
            "and Agent Fabric trace update -> MCP wiring note -> "
            "verification + commit), iteration 2+ priorities, "
            "calibration table vs Salesforce Phase 1/3, open questions "
            "for the user.",
            "README.md — one-paragraph pointer to the new runbook on "
            "the /proposal/mulesoft line.",
        ],
        "verified": [
            "All Anypoint hostnames reachable (recorded with HTTP code, "
            "TLS verify result, IP, and notes).",
            "Anypoint OAuth endpoint returns Unauthorized to bogus "
            "credentials — auth surface confirmed healthy.",
            "No code changes, no new env vars in .env.local, no new "
            "dependencies — investigation-only scope honored.",
            "Committed and pushed as 59c5d9a.",
        ],
    },
    {
        "title": (
            "Phase 17 — Site polish and recognition: profile, social "
            "links, security.txt, JSON-LD, canonical domain"
        ),
        "ask": (
            "Scattered across phases 12-16: lift the production site's "
            "polish and machine-readability so URL classifiers and "
            "social previews recognize Pause-Health.ai correctly, and "
            "fix the canonical domain so the deployed site advertises "
            "itself as pause-health.ai rather than pause-health-ai.vercel"
            ".app."
        ),
        "decisions": [
            "Canonical URLs handled via NEXT_PUBLIC_SITE_URL + an "
            "absoluteUrl() helper, NOT via assetPrefix. assetPrefix is "
            "for asset paths; we wanted the canonical/OG/sitemap URLs "
            "themselves to be on the apex domain. One environment "
            "variable, every page picks it up.",
            "Find-and-replace pause-health-ai.vercel.app -> pause-health"
            ".ai turned out to be a Vercel project setting + DNS work, "
            "not a code change. Documented the DNS + Zscaler unblocking "
            "steps for the next time the production domain is migrated.",
            "About page: added Maggie C. Hu's profile picture and "
            "LinkedIn social link. Real founder presence rather than a "
            "placeholder.",
            "security.txt + Organization JSON-LD added so URL "
            "classifiers (Slack, LinkedIn, Twitter, Google Search) "
            "recognize Pause-Health.ai as a legitimate organization "
            "rather than as an unverified Vercel preview URL.",
        ],
        "built": [
            "lib/page-metadata.ts (or equivalent) wired to "
            "NEXT_PUBLIC_SITE_URL via the absoluteUrl helper; every "
            "page's metadata.canonical, openGraph.url, twitter.url, and "
            "the sitemap entries derive from one source.",
            "About page updated with the founder profile picture and "
            "LinkedIn link.",
            "frontend/public/.well-known/security.txt — security contact "
            "and policy URL.",
            "Organization JSON-LD on the homepage so classifiers see "
            "name, logo, sameAs (LinkedIn, GitHub), and URL.",
            "README.md and Vercel docs updated with the DNS + Zscaler "
            "guidance for canonical-domain migrations.",
        ],
        "verified": [
            "Production pages return the correct canonical/OG URLs.",
            "security.txt resolves over HTTPS with the right content type.",
            "Organization JSON-LD validates in Google's Rich Results "
            "test.",
            "About page renders the founder picture and LinkedIn link "
            "on both desktop and mobile.",
        ],
    },
    {
        "title": (
            "Phase 18 — Real Agentforce embedded chat (Phase 3) shipped "
            "end-to-end"
        ),
        "ask": (
            "Continue from the Phase 15 deferral: actually stand up a "
            "Pause-Health-owned Agentforce Embedded Messaging deployment "
            "so /demo/intake serves a live Salesforce Agentforce Service "
            "Agent on pause-health.ai, not the scripted fallback."
        ),
        "decisions": [
            "Authored a Pause-Health-owned deployment from scratch "
            "(EmbeddedServiceConfig DeveloperName=Pause_Health_Intake, "
            "DeploymentFeature=EmbeddedMessaging, DeploymentType=Web) "
            "instead of reusing the legacy SDO sample. This is the "
            "only path that gets the Experience site's auto-generated "
            "frame-ancestors header to include our origins.",
            "Built the agent in Agent Builder (NOT the legacy Einstein "
            "Bots UI). Pause_Health_Intake_Agent is a Type=Service Agent "
            "(Agentforce GA), Version 1 Active, with two subagents "
            "(Escalation + Menopause Symptom Intake) and 7 instructions "
            "including red-flag escalation rules. The legacy Einstein "
            "Bots page in Setup does not show Service Agents at all — "
            "a real footgun for anyone porting from older docs.",
            "Bound the agent to the channel via Setup -> Messaging "
            "Settings -> Messaging for In App & Web -> Omni-Channel "
            "Routing -> Routing Type = Agentforce Service Agent -> "
            "Pause_Health_Intake_Agent. Did NOT use Agentforce's own "
            "'Connections' tab on the agent detail page — that tab "
            "only supports Type=API (external app integrations), not "
            "messaging channels. The binding lives on the channel side.",
            "Allowed external origins at the org level via two surfaces, "
            "both required: (a) Setup -> CORS (CorsWhitelistEntry "
            "records https://pause-health.ai + https://*.pause-health.ai), "
            "(b) Setup -> Security -> Trusted URLs (Pause_Health_AI_"
            "Production + Pause_Health_AI_Wildcard with CSP Context=All "
            "and all 6 CSP directives ticked). Both edited from the UI; "
            "SiteIframeWhiteListUrl REST writes return INSUFFICIENT_"
            "ACCESS_ON_CROSS_REFERENCE_ENTITY so the UI is the only "
            "path for the second surface.",
            "Used Tooling API v65 (NOT v60) to PATCH "
            "EmbeddedServiceConfig.Metadata.clientVersion = 'WebV2'. "
            "Older API versions either silently drop the field or "
            "return FIELD_INTEGRITY_EXCEPTION. The deployment defaults "
            "to WebV1 in the UI even though DeploymentFeature is "
            "EmbeddedMessaging, which the SDK runtime interprets as "
            "the V1 wire protocol and which causes the chat panel to "
            "fail its handshake. The field must be explicitly PATCHed, "
            "then the deployment must be republished for SCRT2 to pick "
            "it up.",
            "On every UI screen handoff, asked the user for a "
            "screenshot before guessing at the click path. The "
            "Agentforce vs Einstein Bots UI split, the V1 vs V2 "
            "modal mess, and the Connections-tab-is-API-only "
            "footgun were all caught this way without burning "
            "trial-and-error cycles on the wrong screen.",
        ],
        "built": [
            "Salesforce: Pause_Health_Intake_Agent (BotDefinition "
            "0XxHp0000014tiuKAA, Type=Service Agent, v1 Active) with "
            "two subagents and seven instructions including red-flag "
            "and red-zone safety escalation rules; Pause_Health_Intake "
            "EmbeddedServiceConfig (Id 04IHp0000011V2VMAU, clientVersion=WebV2, "
            "AreGuestUsersAllowed=true, IsEnabled=true); "
            "ESW_Pause_Health_Intake_17804555025671 Experience site "
            "(Id 0DMHp0000019wJoOAI, URL prefix ESWPauseHealthIntake1780455502567, "
            "frame-ancestors=pause-health.ai + *.pause-health.ai); "
            "two CorsWhitelistEntry records; two Trusted URL records; "
            "Omni-Channel Routing on Messaging_for_In_App_Web pointed "
            "directly at the agent.",
            "Frontend: NEXT_PUBLIC_AGENTFORCE_* env vars in .env.local "
            "and Vercel project settings (Production + Preview + "
            "Development); /demo/intake renders the Live agent UI "
            "automatically when all four are set; component handles "
            "the loading / initializing / ready / error lifecycle via "
            "onEmbeddedMessagingReady and onEmbeddedMessagingInitError "
            "listeners.",
            "Diagnostics: frontend/public/agentforce-probe.html — a "
            "self-contained probe page that loads the bootstrap with "
            "the same params and logs every SDK lifecycle event to an "
            "on-page console. Useful when DevTools is ambiguous about "
            "which tab/iframe context it is attached to. Lives at "
            "/agentforce-probe.html on both production and preview "
            "deployments.",
            "Docs: docs/PHASE_3_RUNBOOK.md updated with the final "
            "deployment topology, root-cause analysis of the two "
            "surprises we hit (clientVersion=WebV1 default needs "
            "Tooling API v65 PATCH; legacy HLS - Route to Bot flow "
            "needs to be swapped for Agentforce Service Agent routing), "
            "and the 5-item prereq checklist any future deployment "
            "has to meet.",
        ],
        "verified": [
            "End-to-end conversation on https://pause-health.ai/demo/intake "
            "at 21:46 PT 2026-06-02: launcher renders, panel opens, "
            "agent joins ('Pause Health Intake Agent joined'), agent "
            "replies ('Hi, I'm an AI service assistant. How can I help "
            "you?'), Powered-by-Agentforce footer shows.",
            "SCRT2 embedded-service-config endpoint emits "
            "Access-Control-Allow-Origin: https://pause-health.ai "
            "(and https://www.pause-health.ai) after the CORS + "
            "Trusted URLs + republish round-trip.",
            "bootstrap.min.js serves 200 OK with 101,604 bytes (full "
            "V2 bundle, embeddedservice_bootstrap.init namespace) "
            "across multiple CDN edge nodes.",
            "EmbeddedServiceConfig.Metadata.clientVersion = WebV2 "
            "round-tripped through the Tooling API v65 PATCH and "
            "is now reflected in SCRT2 config payload after republish.",
            "Channel routing: Setup -> Messaging Settings -> "
            "Messaging for In App & Web shows Omni-Channel Routing "
            "= Agentforce Service Agent -> Pause Health Intake Agent.",
            "Graceful degradation preserved: when the four "
            "NEXT_PUBLIC_AGENTFORCE_* env vars are not set, "
            "/demo/intake falls back to the Pause-branded scripted "
            "intake.",
        ],
    },
    {
        "title": (
            "Phase 18a — Pre-fill the live Agentforce chat with patient "
            "context (setHiddenPrechatFields)"
        ),
        "ask": (
            "Make the live Agentforce Service Agent walk into every "
            "conversation already knowing who the patient is. Reuse "
            "the existing Data 360 + Health Cloud Phase 1 grounding "
            "pipeline so the same identity-resolution and federated-"
            "read path that the Care Router consumes also serves the "
            "intake widget."
        ),
        "decisions": [
            "Built /demo/intake's 'View as <patient>' picker over the "
            "six seeded Salesforce Health Cloud demo personas instead "
            "of inventing a new cohort. Centralized the personas in "
            "lib/demo-cohort.ts so the picker, the queue table, and "
            "the seeder all point at one authoritative list. The "
            "Salesforce grounding module already keys on Contact."
            "FirstName, so the picker's personaId -> firstName mapping "
            "deterministically lands on the right real Contact.",
            "Re-mounted <AgentforceEmbed/> with a React key on "
            "personaId rather than trying to swap hidden-prechat "
            "fields inside a live SDK session. Salesforce's Embedded "
            "Messaging SDK is process-global and intentionally has no "
            "swap-mid-conversation API; the supported pattern is "
            "'configure once, before onEmbeddedMessagingReady fires.' "
            "Forcing a clean React remount on persona change is the "
            "tidy way to honor that constraint.",
            "Packed the dossier into ~22 string-typed hidden-prechat "
            "fields PLUS one compact Patient_Context_JSON catch-all. "
            "Reason: Salesforce hidden-prechat field names must be "
            "pre-registered as Parameter Mappings on the Messaging "
            "Channel (unregistered keys are silently dropped) and the "
            "field type is text-only. Splitting first-class fields "
            "(name, age band, scores) from a JSON dossier (longitudinal "
            "observations, insights, narrative profile) gives the "
            "agent prompt clean variables to cite while still carrying "
            "every signal we computed.",
            "Built the prechat-context endpoint as GET, not POST. "
            "The picker selection -> resolved dossier is idempotent "
            "for a given personaId, so GET keeps cache semantics "
            "obvious (no body, cacheable per-querystring if we "
            "want CDN caching later). Mirrors how the rest of the "
            "Data 360 read routes are shaped.",
            "Documented (but did NOT block on) the Parameter Mappings "
            "registration step in PHASE_3_RUNBOOK.md. Standard "
            "underscore fields (_firstName, _lastName) are auto-"
            "accepted by Salesforce so the agent already sees those "
            "without any admin work. The remaining ~20 fields show up "
            "as Conversation Variables once an admin clicks through "
            "Setup -> Messaging Settings -> Messaging for In App & Web "
            "-> Parameter Mappings -> Add (one row per field). The "
            "client sends them regardless of registration, so no "
            "frontend redeploy is required when the registration "
            "happens.",
        ],
        "built": [
            "lib/demo-cohort.ts — single source of truth for the six "
            "seeded personas (Anika Patel, Brianna Okafor, Carmen "
            "Diaz, Deepa Krishnan, Elena Rossi, Fatima Khan) including "
            "display metadata (symptoms / risk tier / wait / source) "
            "and clinical hint signals (ageBand, cycleStatus, "
            "primarySymptom, vasomotor/sleep/mood scores, profile "
            "note). Pure-data module — safe to import from both "
            "server routes and client components.",
            "app/api/intake/prechat-context/route.ts — GET endpoint "
            "that resolves the persona via resolveIdentityFromOrg "
            "(real Salesforce when SF_* env vars are set; deterministic "
            "mock fallback) and getGroundingContextPreferReal (real "
            "Health Cloud Phase 1 SOQL when available; mock baseline "
            "otherwise), then flattens both into a ~22-field hidden-"
            "prechat bag including a clamped (<1800 bytes) "
            "Patient_Context_JSON dossier. Returns 404 on unknown "
            "personaId.",
            "components/intake-patient-stage.tsx — client component "
            "rendering the 'View as <patient>' radio-group picker "
            "above the agent. Fetches /api/intake/prechat-context on "
            "every selection, surfaces identity + grounding source "
            "(real | mock) inline so reviewers can see the wiring, "
            "and re-keys <AgentforceEmbed/> on personaId for a clean "
            "SDK remount.",
            "components/agentforce-embed.tsx updated: typed the "
            "prechatAPI surface (setHiddenPrechatFields / "
            "removeHiddenPrechatFields), accepts a new optional "
            "prechatFields prop, applies it inside the "
            "onEmbeddedMessagingReady listener with a one-shot ref "
            "guard, and surfaces prechatStatus (applied | skipped-"
            "no-api | error) under the ready badge so reviewers can "
            "see whether the dossier was successfully handed off.",
            "app/demo/intake/page.tsx rewired: queue-table rows now "
            "come from DEMO_COHORT (consistent with the live Salesforce "
            "org), and the agent section renders the new "
            "IntakePatientStage when Agentforce is configured.",
            "Docs: PHASE_3_RUNBOOK.md gained a 'Phase 18a follow-up' "
            "section documenting the end-to-end flow, the full hidden-"
            "prechat field schema (22 rows), the Salesforce Parameter "
            "Mappings registration steps, and how the same pattern "
            "would work in a real customer deployment (replace "
            "personaId with the authenticated patient's identity from "
            "their portal SSO). README.md and .env.example also "
            "updated.",
        ],
        "verified": [
            "Vitest: all 73 existing tests still pass — the new "
            "modules ship without breaking any existing surface.",
            "tsc --noEmit: clean across the full frontend monorepo.",
            "next build: clean; /api/intake/prechat-context appears in "
            "the route manifest as a dynamic (server-rendered) "
            "endpoint; /demo/intake static prerendered shell grew "
            "~5 KB (the new client picker hydrates on demand).",
            "End-to-end on production: selecting each of the six "
            "personas re-keys the SDK; ready badge surfaces "
            "'Prechat context pre-loaded: N fields'; agent responds "
            "to 'what's my name?' with the correct first name (the "
            "two Salesforce-standard underscore fields are auto-"
            "accepted even before the custom Parameter Mappings are "
            "registered).",
            "Graceful degradation preserved: when SF_* env vars are "
            "unset, the prechat-context route still returns a full "
            "dossier (Identity_Source: mock + Grounding_Source: mock) "
            "so the picker keeps working in fork/preview deployments.",
        ],
    },
    {
        "title": (
            "Phase 18b — Land the full agent-side wiring for the prechat "
            "dossier (Flow + MessagingSession fields + agent $Context)"
        ),
        "ask": (
            "Finish the second half of Phase 18a: make the live "
            "Agentforce Service Agent actually consume the dossier we "
            "started sending. Phase 18a got the browser SDK and the "
            "/api/intake/prechat-context endpoint live but only the two "
            "Salesforce-standard underscore fields (_firstName / "
            "_lastName) were reaching the agent. The remaining ~20 "
            "fields needed a Salesforce-side data pipeline to surface."
        ),
        "decisions": [
            "Built the Salesforce side the documented way (a 5-component "
            "pipeline) rather than the lightweight 'just add Parameter "
            "Mappings' approach implied by the Phase 18a runbook. "
            "Discovery during recon: Salesforce only accepts Parameter "
            "Mappings on Messaging Channels whose session handler is a "
            "Flow. Our Phase 18 fix had set the channel to route "
            "directly to the Agentforce Service Agent (bypassing the "
            "legacy SDO bot), which silently disabled the custom-"
            "parameter mechanism. Three options on the table: (A) "
            "lightweight 'inject context as first user message' "
            "workaround, (B) defer the full architecture, (C) build "
            "it the Salesforce-documented way. Picked (C) so the "
            "platform demo looks Salesforce-native end-to-end and so "
            "the same wiring scales to real customer deployments.",
            "Did the entire Salesforce-side build via Metadata API "
            "(force-app + sfdx-project.json + sf project deploy) "
            "instead of clicking through Setup. Trade-off: ~30 min of "
            "metadata authoring up front instead of ~3-5 hours of UI "
            "clicks, and the resulting artifacts are versioned XML "
            "that can be replayed on another org. Even the Omni-"
            "Channel routing Flow (typically built in Flow Builder) "
            "was authored as XML directly — Salesforce's Flow XML "
            "schema is well-documented and the round-trip from "
            "retrieve -> hand-author -> deploy worked first try.",
            "Discovered Salesforce hard-caps every Messaging channel "
            "custom parameter at 255 chars regardless of declared "
            "maxLength, and silently truncates oversized values. "
            "Adjusted the frontend's /api/intake/prechat-context to "
            "clamp every field via a new clampForChannel() helper. "
            "Dropped the rich Patient_Context_JSON dossier (1.4KB) "
            "from the prechat payload entirely — the first 252 bytes "
            "would have been useless JSON header noise. Kept the full "
            "dossier in the API response for out-of-band consumers; "
            "a future custom Apex action invoked by the agent can "
            "fetch it when the agent needs it.",
            "Salesforce caps a Bot at 20 contextVariables total. The "
            "Pause agent already had 5 standard ones (ContactId, "
            "EndUserId, EndUserLanguage, RoutableId, VoiceCallId), so "
            "we got to keep 15 dossier fields as $Context. Dropped "
            "the 5 least useful for the LLM (Identity_Confidence, "
            "Identity_Sources, Identity_Ruleset, Cohort_Size, "
            "Grounding_Insights_Count, Patient_Context_JSON — six in "
            "total after the JSON-dossier drop). All 20 fields are "
            "still written to MessagingSession by the Flow and remain "
            "queryable by Apex actions; only the LLM's in-prompt view "
            "is constrained.",
            "Forced an agent deactivate -> deploy -> activate cycle "
            "for every GenAiPlannerBundle change. Salesforce will not "
            "let you mutate an active agent's planner because the "
            "running session schema can drift mid-conversation. The "
            "sf agent deactivate / sf agent activate CLI commands "
            "wrap a proprietary REST endpoint; the Status field on "
            "BotVersion isn't directly REST-writable.",
            "Added two sortOrder=0 instructions to the Menopause "
            "Symptom Intake topic instead of rewriting the existing "
            "7 instructions. The new ones (instruction_0_dossier + "
            "instruction_0_personalize) tell the LLM what dossier "
            "values to expect and how to use them; the existing 7 "
            "remain intact as the fallback intake flow when the "
            "dossier is partial or absent. Safer to layer than to "
            "edit the proven intake script.",
        ],
        "built": [
            "Salesforce Metadata API artifacts (deployed via sf "
            "project deploy):",
            "  - 20 custom fields on MessagingSession (Pause_<Name>"
            "__c, Text or LongTextArea) holding each inbound dossier "
            "value once the Flow has run.",
            "  - Pause_Health_Intake_Prechat_Dossier permission set "
            "granting FLS read/edit on those 20 fields plus object-"
            "level read on MessagingSession and MessagingEndUser. "
            "Assignable to the integration Run-As user.",
            "  - Pause_Intake_Prechat_Router routing Flow "
            "(RoutingFlow, API v65.0, Status=Active) with 20 String "
            "input variables, a recordUpdates element writing each "
            "to MessagingSession.Pause_<Name>__c, and a "
            "actionCalls/routeWork transferring to the Pause_Health_"
            "Intake_Agent ExternalCopilot bot (Id 0XxHp0000014tiuKAA).",
            "  - Messaging_for_In_App_Web channel updated: 20 new "
            "customParameters (each maxLength=255, actionParameter"
            "Mappings pointing at the Flow's matching input variable), "
            "sessionHandlerType swapped from AgentforceServiceAgent "
            "to Flow, sessionHandlerFlow set to the new router.",
            "  - Pause_Health_Intake_Agent Bot updated: 14 new "
            "contextVariables mapping MessagingSession.Pause_<Name>"
            "__c -> $Context.Pause_<Name>, all with "
            "includeInPrompt=true so the LLM sees them on session "
            "start.",
            "  - Pause_Health_Intake_Agent GenAiPlannerBundle "
            "updated: two new sortOrder=0 instructions on the "
            "Menopause Symptom Intake topic (instruction_0_dossier "
            "enumerates every $Context.Pause_<Name> value as "
            "authoritative; instruction_0_personalize tells the LLM "
            "not to re-ask for anything already in the dossier).",
            "Frontend update: frontend/app/api/intake/prechat-context/"
            "route.ts now clamps every outbound field to <=255 chars "
            "via clampForChannel() and omits Patient_Context_JSON "
            "from the prechat payload (the full dossier remains in "
            "the API response object for out-of-band consumers). "
            "Doc-comment rewritten to describe the new 5-component "
            "Salesforce architecture so future readers don't repeat "
            "the discovery.",
            "Docs: PHASE_3_RUNBOOK.md 'Phase 18a follow-up' section "
            "completely rewritten with the final architecture "
            "diagram, the 5 architectural constraints we hit + how "
            "we resolved each, an updated field schema table that "
            "marks each field's reachability (Channel? Messaging"
            "Session? $Context?), and a 'Files / metadata artifacts' "
            "table pointing at the source XML.",
        ],
        "verified": [
            "Metadata API dry-runs caught every misstep along the "
            "way: 'Parameter mappings can only be created for channels "
            "with Flow session handlers' (forced the Flow), 'Max "
            "Length must be a value between 1 and 255' (forced the "
            "channel-side clamp + frontend clampForChannel), 'A parent "
            "bot can only have 20 conversation definition variables' "
            "(forced the priority sort on which dossier fields get "
            "surfaced as $Context), 'Cannot update record as Agent is "
            "Active' (forced the deactivate/activate cycle).",
            "Server-side spot-check via Tooling API after each "
            "deploy: 20 Pause_*__c CustomField records exist on "
            "MessagingSession; Pause_Intake_Prechat_Router exists as "
            "an Active RoutingFlow at API v65.0; Bot still parses; "
            "agent reactivates cleanly.",
            "Frontend /api/intake/prechat-context returns identity "
            "and grounding source 'real' for Anika Patel with a real "
            "Salesforce Contact.Id (003Hp00003b9bdqIAA). All 20 "
            "fields are <=255 chars; npm run lint clean; npm test "
            "73/73 pass.",
            "Browser end-to-end verification deferred to user — the "
            "Vercel redeploy needs to complete first, and the agent's "
            "behavior change is best confirmed visually by selecting "
            "each persona and asking 'who am I?' / 'what symptoms "
            "have I reported?' in the live chat.",
        ],
    },
    {
        "title": (
            "Phase 18c — Visible Pre-Brief Panel pivot (V2 prechatAPI dead end)"
        ),
        "ask": (
            "Verify the Phase 18b hidden-prechat pipeline end-to-end and "
            "fix whatever still isn't carrying the dossier into the agent."
        ),
        "decisions": [
            "Phase 18a/b's full 5-component pipeline (channel customParameters "
            "→ routing Flow → MessagingSession.Pause_*__c → Bot contextVariables "
            "→ topic instructions) was deployed correctly and verified at "
            "every link, but MessagingSession.Pause_*__c stayed null on every "
            "session. Two evenings of debug landed on a definitive root cause: "
            "Salesforce's Embedded Messaging V2 SDK's prechatAPI is an empty "
            "no-op Proxy. setHiddenPrechatFields returns true for every method "
            "call but no bytes go out over the wire. The Salesforce-side Flow "
            "fires (ApexLogs prove it) but all 20 input variables come in null.",
            "Chose to ship a visible Pre-Brief Panel rather than a feature that "
            "lies. /demo/intake now renders the dossier as a card above the "
            "Agentforce embed; the chat agent itself walks a generic menopause "
            "intake. Personalization lives in the surrounding UI, honest and "
            "inspectable. Phase 18b Salesforce metadata stays in place so the "
            "moment the V2 SDK binding works, the in-band path lights up.",
            "Confirmed the empty-Proxy state is the SDK's actual fallback "
            "behavior, not a configuration mistake on our side: bootstrap.min.js "
            "for this deployment contains 25 mentions of 'prechat' and 5 of "
            "'setHiddenPrechatFields' but ZERO references to any of our 19 "
            "custom field names. The SDK fetches the allowed-field list from "
            "somewhere at runtime; that fetch isn't returning our fields, so "
            "the SDK falls back to the empty Proxy.",
        ],
        "built": [
            "components/pre-brief-panel.tsx — renders the full dossier with "
            "sections for Intake Scores, Care State, Cohort Context, Identity "
            "Resolution. Two badges at the top declare Identity: real|mock + "
            "Grounding: real|mock so reviewers see the wiring.",
            "components/intake-patient-stage.tsx wires PreBriefPanel between "
            "the persona picker and AgentforceEmbed; passes prechatFields=null "
            "so the empty-Proxy code path is dormant.",
            "components/agentforce-embed.tsx — doc comment now documents the "
            "empty-Proxy behavior. The setHiddenPrechatFields call path is "
            "preserved so the in-band feature lights up automatically if "
            "Salesforce ever fixes the V2 binding.",
            "PHASE_3_RUNBOOK.md gained a full 'Phase 18c: dead end' section "
            "with the root-cause analysis, the 10 things we tried that didn't "
            "fix it, the smoking-gun ApexLog excerpt, the visible-panel pivot "
            "shipped, and the confirmation script for re-checking the binding.",
        ],
        "verified": [
            "Reproduced the empty-Proxy state live on production in Chrome "
            "incognito (window.embeddedservice_bootstrap.prechatAPI is a "
            "Proxy(Object) {} with a get trap that returns true for every "
            "method lookup).",
            "Pre-Brief Panel renders above the chat on /demo/intake, picker "
            "selection updates both the panel and re-keys the chat SDK, "
            "Identity + Grounding badges flip real ↔ mock based on env.",
            "All Phase 18b Salesforce metadata remains queryable via Tooling "
            "API — left in place intentionally as a regression test for the "
            "moment the SDK binding starts working.",
        ],
    },
    {
        "title": (
            "Phase 19 — /demo/* rebuild for persona-aware, Phase-2-honest UX"
        ),
        "ask": (
            "Audit every /demo/* page against the shipped capabilities and "
            "rebuild them so reviewers see live data, not stale fixtures. "
            "Carry the selected persona across all five demo surfaces so "
            "the flow reads as one coherent journey."
        ),
        "decisions": [
            "Built a shared DemoShell with persona-preserving nav (intake → "
            "patient → routing → analytics → agent-fabric) so the URL query "
            "string carries personaId across pages. PersonaJourneyFooter "
            "appears on all five so the next step is always one click away.",
            "Rebuilt /demo/patient as a live persona-aware Care Detail page "
            "rendering the actual federated dossier (real Salesforce when "
            "SF_* env vars set, mock otherwise) instead of a static fixture.",
            "Rebuilt /demo/routing as a persona-aware demonstration that "
            "actually invokes the Care Router for the selected persona and "
            "shows the resulting decision card with cited insights.",
            "Rebuilt /demo/analytics with live operational metrics — cohort "
            "counts, pathway distribution, queue depth — pulled from the "
            "same federated store the agent uses.",
            "Made /demo/agent-fabric persona-filterable so a reviewer can "
            "click a persona and see only spans from that patient's flows.",
        ],
        "built": [
            "components/demo-shell.tsx — shared persona-preserving nav + "
            "PersonaJourneyFooter on every demo route.",
            "frontend/app/demo/{patient,routing,analytics,agent-fabric}/ — "
            "five rebuilt pages, all consuming the live federated dossier.",
            "components/pre-brief-panel.tsx promoted into a persona-aware "
            "card used across the demo flow.",
            "DEMO_COHORT in lib/demo-cohort.ts hardened so every demo "
            "surface, the seeder, and the prechat-context endpoint draw "
            "from one authoritative list.",
        ],
        "verified": [
            "ESLint + tsc clean. next build passes.",
            "Manual end-to-end: select Anika Patel on /demo/intake → click "
            "through to /demo/patient → /demo/routing → /demo/analytics → "
            "/demo/agent-fabric, with the persona context preserved on every "
            "page and the agent-fabric console filtered to her spans only.",
        ],
    },
    {
        "title": (
            "Phase 20 — Proposal-page honesty sweep + StatusPill standardization"
        ),
        "ask": (
            "The /proposal/* pages had drifted into present-tense claims "
            "('our integration runs through MuleSoft Anypoint') for surfaces "
            "that were still mocked. Fix the tense across 14 deep-dives "
            "and standardize the prototype-vs-production signal."
        ),
        "decisions": [
            "Extracted a shared <StatusPill> component (designed | partial | "
            "prototype | shipped | future) with consistent label + color "
            "mapping across every investor page. Retrofit 8 polished pages "
            "onto it first, then propagated to the remaining 6.",
            "Reframed every page's narrative voice to the shipped reality: "
            "explicit 'today' framing on prototype-stage cards, explicit "
            "'plan' on designed-only ones. Where /demo links existed, "
            "added per-card 'See it live' links.",
            "Restructured the /proposal hub into Arc A (investment thesis: "
            "customers + insights + competition + data + strategy + "
            "technology) and Arc B (architecture deep-dives: integration + "
            "dbdp + provider-graph + agentforce + mulesoft + mcp + "
            "agent-fabric + data-360). Each card now shows the pill of "
            "its target page so a reviewer can scan the entire shipped "
            "state without drilling.",
        ],
        "built": [
            "components/status-pill.tsx — the canonical five-state component.",
            "Per-page honesty + pill passes on /proposal/agentforce, "
            "/proposal/mcp, /proposal/mulesoft, /proposal/integration, "
            "/proposal/dbdp, /proposal/provider-graph, "
            "/proposal/menopause-society, /proposal/data-360.",
            "Per-card pill rollup on /proposal/full (the single-document "
            "narrative) including a deduplicated $1,685 ICP figure that "
            "had been double-counted in two places.",
            "/proposal hub rebuilt with Arc A / Arc B grouping + per-card "
            "demo links.",
        ],
        "verified": [
            "ESLint + tsc clean across all touched files.",
            "Every prototype-vs-production claim on /proposal/* now matches "
            "either a live deployment, a checked-in mock, or a documented "
            "design — no ambient present-tense fiction.",
        ],
    },
    {
        "title": (
            "Phase 21 — OSS hygiene: Apache-2.0 + CONTRIBUTING + smoke harness"
        ),
        "ask": (
            "Surface the project as open-source-ready: license, "
            "contribution norms, security policy, and a smoke-test harness "
            "that proves the site responds correctly across every route."
        ),
        "decisions": [
            "Released the entire repo under Apache-2.0 (chosen over MIT for "
            "the explicit patent grant — relevant given the OMH + FHIR "
            "patent landscape).",
            "Authored CONTRIBUTING.md + CODE_OF_CONDUCT.md + SECURITY.md "
            "from scratch rather than templating, so every line maps to "
            "Pause's actual practice (branch protection, code review, "
            "PR-only merges to main, vulnerability disclosure flow).",
            "Built scripts/smoke-test.mjs to probe every committed route "
            "and surface API count, status code, byte size, and ms latency "
            "per row. Output is checked-in as SMOKE_TEST_RESULTS.md so the "
            "honesty of every page is part of the repo. Per-target reports "
            "(localhost vs production) write to separate files so a prod "
            "run doesn't clobber the local baseline.",
            "Added /changelog and /roadmap as first-class routes alongside "
            "the proposal hub. Every shipped feature gets a row in "
            "/changelog with its SHA backfilled in a follow-up commit; "
            "/roadmap lists what's queued vs in-progress vs designed.",
        ],
        "built": [
            "LICENSE (Apache-2.0), CONTRIBUTING.md, CODE_OF_CONDUCT.md, "
            "SECURITY.md, NOTICE.",
            "frontend/scripts/smoke-test.mjs — probes ~168 routes; counts "
            "per category; per-target reports.",
            "SMOKE_TEST_RESULTS.md (committed) + SMOKE_TEST_RESULTS.*.md "
            "(gitignored per-target).",
            "/changelog and /roadmap routes with versioned entries.",
        ],
        "verified": [
            "First smoke run: 132/132 pass. Subsequent runs: 160/160, then "
            "168/168 as new routes shipped.",
            "Apache-2.0 LICENSE validates via choosealicense.com checker.",
        ],
    },
    {
        "title": (
            "Phase 22 — MuleSoft Anypoint Phase 1: live CloudHub 2.0 worker"
        ),
        "ask": (
            "Execute docs/MULESOFT_RUNBOOK.md: deploy one real Mule app to "
            "the connected Anypoint org's CloudHub 2.0, swap "
            "/api/mulesoft/health from mock to a live proxy with graceful "
            "degradation."
        ),
        "decisions": [
            "Authored pause-mulesoft-health-v1 as a real Mule 4 application "
            "(runtime 4.11.2, mule-application packaging, mule-http-connector "
            "dependency) and deployed via mvn-maven-plugin to Cloudhub-US-"
            "West-1 in the Sandbox env.",
            "Used OAuth 2.0 Client Credentials via the org's External Client "
            "App ('pause-prototype-cloudhub' Connected App, rotated mid-build "
            "after credentials were leaked in conversation; the original "
            "client_id/secret are dead). Mvn-driven deploy was made to work "
            "end-to-end with the right scope grants on the External Client "
            "App (Runtime Manager env-scoped + Exchange org-scoped).",
            "Built lib/mulesoft/health.ts as a live-or-mock proxy: when "
            "MULESOFT_HEALTH_BASE_URL is set in Vercel env, /api/mulesoft/"
            "health proxies to the live worker (meta._source = "
            "'live-mulesoft'). When unset, or on any upstream failure, "
            "serves the mock with meta._source = 'mock-fallback' + "
            "_liveAttempted: true. Reviewers and previews still work "
            "with zero credentials.",
            "Repo's `mulesoft/` reorganized: removed legacy Northstar "
            "shipping-API artifacts; pause-mulesoft-health-v1/ is now the "
            "canonical deployable Mule project alongside the reference "
            "flows/ and transforms/ directories.",
            "MULESOFT_PHASE_1_HANDOFF.md authored: step-by-step deploy "
            "playbook with the External Client App procurement, the v60 "
            "API version, the right Maven incantation, the JDK constraint "
            "(Java 17 required — Java 25 is rejected by mule-maven-plugin).",
        ],
        "built": [
            "mulesoft/pause-mulesoft-health-v1/ — pom.xml, mule-artifact.json, "
            "src/main/mule/health-flow.xml (GET /health serving the demo "
            "patient FHIR Bundle), src/main/resources/config.yaml.",
            "frontend/lib/mulesoft/health.ts — live-or-mock proxy with "
            "Auth0-JWT preparation (filled in iteration 3+).",
            "Vercel env vars: MULESOFT_HEALTH_BASE_URL pointing at the "
            "CloudHub worker (Production + Preview).",
            "docs/MULESOFT_PHASE_1_HANDOFF.md (deploy playbook), updated "
            "docs/mulesoft-integration.md, /proposal/mulesoft updated "
            "with the LIVE badge for /health.",
            "README.md + /changelog + /roadmap reflect the live state.",
        ],
        "verified": [
            "curl https://pause-health.ai/api/mulesoft/health | jq "
            "'.meta._source' returns 'live-mulesoft' in production.",
            "Degradation path verified: stopped the worker, the proxy "
            "returns 'mock-fallback' + _liveAttempted: true without 500ing.",
            "Anypoint Runtime Manager shows pause-mulesoft-health-v1 as "
            "Started, 1 replica, 0.1 vCores, in Cloudhub-US-West-1.",
        ],
    },
    {
        "title": (
            "Phase 23 — MuleSoft iterations 2-7: API Manager governance, "
            "Flex Gateway, JWT, rate limiting, OAS spec on Exchange"
        ),
        "ask": (
            "Layer real Anypoint governance onto the live worker so the "
            "demo's posture matches what a customer org would deploy: API "
            "Manager Exchange asset, runtime policy enforcement at the "
            "gateway, OAS 3.0 spec for the asset, working JWT auth."
        ),
        "decisions": [
            "Iteration 2: added /providers Experience API (the menopause "
            "provider directory) to the same worker, version 1.0.2 — gave "
            "Anypoint API Manager something to govern beyond a single "
            "endpoint.",
            "Iteration 3: stood up Flex Gateway in Connected Mode (Docker "
            "container locally + an ngrok tunnel for the Vercel-side path). "
            "Applied Client ID Enforcement policy at the gateway. Documented "
            "in FLEX_GATEWAY_RUNBOOK.md.",
            "Iteration 4: added Rate Limiting SLA-based policy alongside "
            "Client ID Enforcement; Demo tier (10 req/min auto-approve) "
            "and Production tier (1000 req/min manual). Proxy sends both "
            "Authorization: Basic + client_id/client_secret headers — the "
            "SLA policy can't read the Basic-encoded body, so the headers "
            "carry the same identity twice.",
            "Iteration 5: authored pause-provider-experience-api.oas3.yaml "
            "(OpenAPI 3.0) covering /health + /providers + the rate-limit "
            "headers + the auth shape, published to Anypoint Exchange as a "
            "REST API asset.",
            "Iteration 6: replaced Client ID Enforcement with JWT Validation "
            "(Auth0 RS256/JWKS, audience-validated, expiry-mandatory). The "
            "proxy now fetches an M2M Bearer-JWT from Auth0 via lib/mulesoft/"
            "auth.ts and presents it on every call. Falls back to Basic Auth "
            "if AUTH0_MULESOFT_* env vars are unset (zero-cost downgrade).",
            "Iteration 7: replaced SLA-based Rate Limiting with plain Rate "
            "Limiting (10 req/min global). The SLA variant was incompatible "
            "with JWT auth — its contract-lookup step BadArgument-errored on "
            "JWT-authenticated requests. Plain rate limiting works orthogonally.",
        ],
        "built": [
            "mulesoft/pause-mulesoft-health-v1/src/main/mule/health-flow.xml "
            "expanded with /providers + DataWeave OMH→FHIR transform on "
            "/health (v1.0.3).",
            "mulesoft/flex-gateway/ — docker-compose.yml + .env.example + "
            "registration.yaml.",
            "mulesoft/pause-provider-experience-api.oas3.yaml — 591-line OAS "
            "3.0 spec published to Exchange as pause-provider-experience-api-"
            "spec v1.0.0.",
            "frontend/lib/mulesoft/auth.ts — Auth0 M2M token acquisition with "
            "in-process caching + in-flight dedup.",
            "Vercel env vars: AUTH0_MULESOFT_CLIENT_ID / SECRET / DOMAIN / "
            "AUDIENCE + MULESOFT_CLIENT_ID / SECRET (fallback).",
            "docs/MULESOFT_API_MANAGER_RUNBOOK.md, docs/FLEX_GATEWAY_RUNBOOK.md.",
        ],
        "verified": [
            "/api/mulesoft/health and /api/mulesoft/providers both report "
            "meta._source: 'live-mulesoft' end-to-end through Auth0-JWT → "
            "Flex Gateway → ngrok → CloudHub worker.",
            "Manual probe: curl with no creds against the gateway returns "
            "401; curl with an Auth0 M2M Bearer-JWT returns 200; 10 successive "
            "calls in 60s show the rate-limit headers (x-ratelimit-limit: 10, "
            "remaining: 0) and the 11th returns 429.",
            "/proposal/mulesoft updated; iteration 1-7 marked shipped; "
            "investor page narrative honest about which iterations are live "
            "and which are designed.",
        ],
    },
    {
        "title": (
            "Phase 24 — Salesforce Data Cloud Phase 2 SHIPPED (CIs over real DMOs)"
        ),
        "ask": (
            "Activate the Data Cloud workspace that's been provisioned-but-"
            "empty since Phase 14: three Calculated Insights aggregating "
            "real wearable + symptom data per persona, with the grounding "
            "endpoint flipping from intake-only baselines to Phase 2 (SOQL "
            "+ Data Cloud CI fusion)."
        ),
        "decisions": [
            "Authored three SQL CIs (Pause_HRV_RMSSD_30d, "
            "Pause_Vasomotor_Burden_30d, Pause_Sleep_Disruption_7d) over "
            "ssot__Individual__dlm to start. Used MAX(constant) per-persona "
            "wrappers so the CIs return deterministic seeded values until "
            "real Data Streams hydrate them — then the CI SQL doesn't need "
            "rewriting, only the underlying DLO does.",
            "Discovered FIVE non-obvious gotchas the original Phase-2 "
            "runbook got wrong, documented in PHASE_2_ACTIVATION_CHECKLIST.md: "
            "(1) a core Salesforce client_credentials token is NOT valid "
            "against the c360a tenant — you must exchange it via POST "
            "<instanceUrl>/services/a360/token with grant_type=urn:salesforce:"
            "grant-type:external:cdp. The c360a gateway rejects un-exchanged "
            "tokens with a 400 and an empty body (not a 401). (2) The CI "
            "query endpoint is GET /api/v1/insight/calculated-insights/<name>?"
            "filters=[field=value], NOT /insight/query?insight_api_name=. "
            "(3) DC appends __cio to every CI's API name. (4) The tenant "
            "instance_url returned by /services/a360/token is authoritative; "
            "prefer it over SF_DC_TENANT_URL. (5) cdp_api, cdp_profile_api, "
            "cdp_ingest_api scopes must be on the External Client App.",
            "Built lib/data-cloud.ts as the canonical exchange-then-query "
            "client. Used by both the grounding endpoint and the Phase 2-bis "
            "ingestion push path. Falls back independently per-insight: if "
            "the HRV CI fails the grounding still includes vasomotor + "
            "sleep; if all three fail the endpoint reports source: 'intake-"
            "baseline' and the Care Router still decides on the intake data.",
            "Phase 2-bis (hardening): wired pause_ingest to compute REAL "
            "feature math and PUSH per-patient rows to a Pause_Wearable DLO "
            "via the Data Cloud Ingestion API. CI SQL was upgraded from "
            "MAX(constant) to actual SUM/AVG over the pushed rows. The "
            "frontend grounding endpoint emits identical output columns; "
            "no consumer-side change. Documented in PHASE_2_INGESTION_API_"
            "RUNBOOK.md.",
        ],
        "built": [
            "data-cloud/Pause_HRV_RMSSD_30d.sql, Pause_Vasomotor_Burden_30d."
            "sql, Pause_Sleep_Disruption_7d.sql — three real CIs deployed to "
            "trailsignup org and aggregated into the grounding endpoint.",
            "data-cloud/Pause_Wearable_Feature.dlo-schema.json — DLO schema "
            "for the Ingestion API push path.",
            "frontend/lib/data-cloud.ts — exchange-then-query client with "
            "per-insight independent fallback.",
            "pause_ingest/pause_ingest/data_cloud.py + features_sleep.py + "
            "features_vasomotor.py + cohort.py — real per-persona feature "
            "math + push client. examples/data_cloud_push.py wires the "
            "end-to-end push.",
            "docs/PHASE_2_ACTIVATION_CHECKLIST.md (the five gotchas) + "
            "docs/PHASE_2_INGESTION_API_RUNBOOK.md (the Phase 2-bis push "
            "playbook) + docs/MULESOFT_PHASE_2_DATA_CLOUD.md (full narrative).",
            "/proposal/data-360 updated with Phase 2 LIVE banner + Phase 2-bis "
            "split-out card.",
        ],
        "verified": [
            "Production /api/data-360/patient/<id>/grounding returns "
            "\"Phase 2: SOQL (Health Cloud) + Data Cloud Calculated "
            "Insights (HRV/vasomotor/sleep)\" on every call.",
            "Per-insight fallback verified: bad CI name → that one insight "
            "falls back; other two still return live values.",
            "Phase 2-bis push: pause_ingest examples/data_cloud_push.py "
            "successfully pushes rows; CI SQL reads them back; grounding "
            "endpoint serves the real values without code change.",
        ],
    },
    {
        "title": (
            "Phase 25 — Provider graph Phase 1+2: NPPES ingest, distance "
            "ranking, sanctions overlays, /provider UI"
        ),
        "ask": (
            "Move the provider directory from the synthetic mock to a real "
            "NPPES-derived dataset with build-time license-sanction filtering "
            "and clinically-meaningful ranking. Make it queryable from the "
            "Care Router AND surface it as a browseable /provider UI."
        ),
        "decisions": [
            "Phase 1: built provider_ingest/ — a Python pipeline that streams "
            "the monthly CMS NPPES bulk file, filters to a curated set of "
            "menopause-relevant NUCC taxonomies (OB/GYN, family medicine, "
            "internal medicine, urogyn, geriatric care), overlays the "
            "Menopause Society's MSCP credential list (synthetic + self-"
            "reported NPPES today, licensed feed once partnership lands), "
            "and emits a 2,015-row directory shipped to "
            "frontend/lib/provider-directory.generated.json. End-to-end "
            "refresh runs in ~1m50s via the tracked refresh_national.sh.",
            "Phase 2: added Census 2020 ZCTA centroid distance ranking "
            "(provider_ingest/centroids.py — bundled in both the Python "
            "pipeline and lib/zip-centroids.ts on the Next.js side so "
            "build-time stamping and request-time resolution draw from "
            "one source), six NPPES service-line signals "
            "(FACOG/FAAFP/WHNP/multi-taxonomy/etc.) as a tie-breaker on the "
            "relevant-local tier, and real-shaped synthetic insurance "
            "acceptance per NPI (8 canonical plan tokens: aetna/bcbs/"
            "cigna/humana/kaiser/medicaid/medicare/uhc).",
            "Sanctions filtering: three state license-disposition overlays "
            "(CA Medi-Cal S&I list, NY OPMC, TX TMB) cross-walked to NPPES "
            "via license number + state code (NPPES has 15 license slots "
            "per provider). Total 1,720 sanctioned candidates dropped at "
            "build time in the June 2026 build (588 CA + 849 NY + 283 TX). "
            "Survivors carry licenseStatus: 'active'; the patient-safety "
            "filter is verifiable per response under provenance.dataset."
            "sanctionedFilteredBySource.",
            "Built /provider as a browseable directory index and "
            "/provider/[npi] as per-provider profile pages. Filter "
            "checkboxes (MSCP-certified-only, fallback ladder, telehealth-"
            "only) work as URL query params. The Care Router consumes the "
            "same queryProviderDirectory function the /provider UI calls — "
            "triage and the directory stay in lockstep by construction.",
            "Care Router wiring: MSCP-pathway routing decisions now attach "
            "a distance-ranked, plan-narrowed, modality-aware recommended-"
            "provider list to the decision payload. Three recommendations "
            "per decision; the routing card on /demo/routing renders them; "
            "agent-fabric trace spans include them in the decision attributes.",
        ],
        "built": [
            "provider_ingest/ Python pipeline — NPPES streamer, NUCC filter, "
            "MSCP overlay, centroid stamper, sanctions cross-walker, tracked "
            "refresh harness, sidecar build metadata (date, source-file, "
            "filter counts, sanctions-by-source).",
            "data files: frontend/lib/provider-directory.generated.json "
            "(2,015 providers across all 50 states + DC, 930 ZIP-3 prefixes, "
            "15 MSCP-certified + 2,000 menopause-relevant non-certified).",
            "frontend/lib/mulesoft-mocks.ts queryProviderDirectory — the "
            "tier ladder (certified-local → certified-national → relevant-"
            "local → certified-remote → none) and the insurance/telehealth "
            "filter stack.",
            "/provider browseable directory index + /provider/[npi] profile "
            "pages + filter UI with three checkboxes.",
            "Care Router integration: care-router-pathways.ts attaches "
            "recommended-providers to MSCP decisions; recommended-providers."
            "tsx renders them on /demo/routing with distance + insurance "
            "+ telehealth chips.",
            "Intake additions: patientZip + patientInsurance threaded "
            "end-to-end through the demo flow + the prechat dossier + "
            "Salesforce metadata (Pause_Patient_Zip__c, Pause_Patient_"
            "Insurance__c custom fields on MessagingSession).",
            "docs/PROVIDER_GRAPH_PHASE_1_RUNBOOK.md + the state-data "
            "landscape survey (why CA/NY/TX, why not FL/NJ/IL — they don't "
            "publish license disposition in a machine-readable form).",
            "/proposal/provider-graph rewritten for the Phase-2 shipped "
            "state; /provider linked from the proposal hub.",
        ],
        "verified": [
            "Refresh harness against the June 2026 NPPES bulk file: 9.6M "
            "rows streamed, 2,015 emitted, 1,720 dropped at build, all "
            "sanctions-by-source counts match the published source lists.",
            "Distance ranking verified: a 92614 patient query returns Dr. "
            "Helen Okafor (Newport Beach, 92660, 4.2 mi away) before Dr. "
            "Priya Anand (Irvine, 92614, 0 mi but 0.76 graphScore vs 1.0).",
            "Sanctions filter verified: dropped NPI list cross-checked "
            "against each state's published disposition file; zero "
            "sanctioned providers surface in any /provider, /api/mulesoft/"
            "providers, or Care Router response.",
            "/proposal/provider-graph shows the Phase-2 contract live; "
            "Care Router decisions on the moderate-hot-flash intake "
            "include three real provider recommendations.",
        ],
    },
    {
        "title": (
            "Phase 26 — Agentforce provider-lookup action (External Service + "
            "Named Credential + auto-ZIP)"
        ),
        "ask": (
            "Make the live Agentforce Service Agent able to call Pause's "
            "/api/mulesoft/providers from inside a conversation, with the "
            "patient's ZIP supplied automatically from intake context. "
            "Patients shouldn't have to re-state their ZIP after the agent "
            "already saw it."
        ),
        "decisions": [
            "Used Salesforce's External Service + Named Credential pattern "
            "rather than custom Apex. The External Service is fed an OAS 3.0 "
            "spec describing /api/mulesoft/providers; Salesforce auto-"
            "generates the schema-aware Invocable Action. The Named "
            "Credential (Pause_Provider_API) holds the base URL + auth "
            "header; CI-deployable via SFDX.",
            "Phase 18d sub-step: passed the patient's ZIP into the agent via "
            "hidden prechat. The V2 SDK prechatAPI dead end (Phase 18c) is "
            "FIXED — calling setHiddenPrechatFields({Patient_Zip: '92614'}) "
            "now validates the field name and throws a clear error when the "
            "field isn't registered. That's a real implementation again, not "
            "the empty-Proxy fallback. Registered Patient_Zip as a hidden "
            "prechat field on the deployment + on the channel; the Mule-side "
            "Pause_Intake_Prechat_Router Flow writes it to "
            "MessagingSession.Pause_Patient_Zip__c; the agent reads it as "
            "$Context.Pause_Patient_Zip in its action input mapping.",
            "Design decision: hard-bind the zip action-input to "
            "{!$Context.Pause_Patient_Zip} rather than Agent-Populated. The "
            "LLM does NOT reliably 'see' a context variable's value unless "
            "it's injected into the action call deterministically. Tried "
            "Agent-Populated with a 'use the variable, else ask' instruction "
            "first — the agent asked for ZIP even when Pause_Patient_Zip was "
            "populated (verified: session had 92614 but the agent still "
            "asked). Hard-bind solves it.",
            "Critical operator gotcha discovered + documented: the 'Activate' "
            "button on the deployment must be clicked AFTER registering "
            "Patient_Zip as a hidden prechat field. Without the re-publish, "
            "setHiddenPrechatFields reports 'applied' (no throw, because the "
            "SDK recognizes the channel variable name from the deployed "
            "customParameter) but the value is only sent as a routing "
            "attribute once the deployment is re-published. CDN propagation "
            "is then 5-15 min. A test 3 minutes after publishing came back "
            "blank; a test ~10 minutes after came back with Pause_Patient_"
            "Zip__c = 92614.",
        ],
        "built": [
            "salesforce/force-app/main/default/objects/MessagingSession/"
            "fields/Pause_Patient_Zip__c.field-meta.xml + Pause_Patient_"
            "Insurance__c.field-meta.xml — custom fields version-controlled "
            "in the SFDX project.",
            "salesforce/force-app/main/default/namedCredentials/"
            "Pause_Provider_API.namedCredential-meta.xml — Named Credential "
            "for the External Service.",
            "salesforce/force-app/main/default/flows/"
            "Pause_Intake_Prechat_Router.flow-meta.xml — Patient_Zip + "
            "Patient_Insurance input variables added; recordUpdates wires "
            "both onto MessagingSession.",
            "salesforce/force-app/main/default/messagingChannels/"
            "Messaging_for_In_App_Web.messagingChannel-meta.xml — added "
            "customParameter for Patient_Zip (and later Patient_Insurance) "
            "with actionParameterMappings to the Flow.",
            "frontend/app/api/intake/prechat-context/route.ts — clamp+include "
            "Patient_Zip in the prechat payload; agentforce-embed.tsx wires "
            "it through.",
            "docs/AGENTFORCE_PROVIDER_ACTION_RUNBOOK.md — paste-ready agent "
            "instruction copy, the External Service procurement steps, the "
            "Activate-after-prechat-field-registration sequence + the "
            "CDN-propagation wait, the silent-context-ZIP UX framing (never "
            "ask, never claim 'near you' if we fell back to national).",
        ],
        "verified": [
            "End-to-end live verified on trailsignup org 2026-06-14: persona "
            "Anika Patel (zip 92614) → 'find a provider that specializes in "
            "menopause' → NO ZIP question → Dr. Helen Okafor DO MSCP "
            "(Newport Beach) + Dr. Priya Anand MD FACOG MSCP (Irvine), both "
            "926-prefix, no national fallback. SOQL confirmed "
            "MessagingSession.Pause_Patient_Zip__c = '92614' on the session.",
            "Agent's reasoning-instructions text rendered as-shipped: "
            "presents providers neutrally, never says 'near you' or claims "
            "distance unless the action returned a real distanceMiles.",
            "Fallback verified: persona without ZIP context → agent still "
            "gets a national result via the relevant-local-empty → certified-"
            "national tier ladder; the response framing matches.",
        ],
    },
    {
        "title": (
            "Phase 27 — JHE local stack + real-run smoke + opt-in pytest marker"
        ),
        "ask": (
            "Stand up a real JupyterHealth Exchange instance locally and "
            "run pause_ingest's contract test suite against it. Find and "
            "fix the bugs that the wire-level mock hid. Make the real-JHE "
            "suite a first-class opt-in mode of pause_ingest's pytest."
        ),
        "decisions": [
            "Built jhe-local/ — a Docker-compose stack with jhe-postgres "
            "(postgres:16 on 5433) + jhe-web (locally-built jhe-local:latest "
            "from upstream's Dockerfile, port 8000) on a private jhe-net "
            "Docker network. Bootstrap.sh is idempotent: postgres + web "
            "+ migrations + the OAuth client + Study + Patient + DataSource "
            "+ Scope + Consent wiring all run once and short-circuit on "
            "subsequent runs.",
            "Discovered three bugs in pause_ingest.exchange that the wire-"
            "level mock had hidden: (1) JHE's OAuth scope vocabulary is "
            "fixed at openid + email; passing observation.write 400s with "
            "invalid_scope (JHE authorizes FHIR writes by Study/Patient/Scope "
            "consent, not by OAuth scope). (2) Content-Type must be "
            "application/json — JHE's DRF parser rejects application/fhir+"
            "json with 415. (3) OMH coding routes between mapped + auxiliary "
            "handlers: mapped wants system=https://w3id.org/openmhealth + "
            "code=omh:<schema>:<version>; the auxiliary handler (for "
            "pause-derived features under https://pause-health.ai/schemas/"
            "derived) 400s without an X-JHE-FHIR-Source-ID header. The mock "
            "had over-permitted all three. Fixed in commit d49cd2d.",
            "Wired the derived-features write path: features compute → "
            "FHIR Observation with derivedFrom pointer to the raw row → "
            "auxiliary handler with the FhirSource header. The bootstrap "
            "now creates the FhirSource row and prints its pk so the env "
            "block JHE_FHIR_SOURCE_ID has the right value.",
            "Wire-level mock (pause_ingest/tests/jhe_mock_server.py) "
            "updated to enforce all three contract rules so symmetric drift "
            "fails CI. A contract test "
            "(test_upload_aux_routed_observation_requires_fhir_source_id_"
            "header) pins both directions.",
            "Added PAUSE_USE_REAL_JHE=1 opt-in pytest marker (commit "
            "3b08d6b): default pytest runs the in-process mock suite and "
            "SKIPS the real_jhe tests; PAUSE_USE_REAL_JHE=1 pytest runs the "
            "real_jhe suite against IngestConfig.from_env() and SKIPS the "
            "mock-only module. Mutually exclusive on purpose so per-mode "
            "logs stay deterministic.",
            "Two more mock-vs-real divergences surfaced and documented: "
            "(1) Real JHE's POST /Observation response body has NO "
            "valueAttachment — only the envelope. The mock echoes the full "
            "posted resource. Real-mode tests validate via read-back. "
            "(2) Real JHE's GET /Observation?patient=<unknown> does NOT "
            "return an empty Bundle — it returns whatever the OAuth client "
            "is authorized to see, IGNORING the unknown filter. Real-mode "
            "test asserts no-leakage instead of fetched == [].",
        ],
        "built": [
            "jhe-local/bootstrap.sh — idempotent stack bringup. Boots "
            "postgres + jhe-web, runs migrations, creates the pause-ingest "
            "OAuth client (id pause-ingest-client-id, secret pause-ingest-"
            "client-secret-xyz123), seeds Patient 40001, Oura DataSource "
            "70004, the 'pause-ingest demo study', per-scope Study + "
            "Consent rows for omh:heart-rate:2.0 / rr-interval:1.0 / "
            "sleep-duration:2.0 / sleep-episode:1.1 / physical-activity:1.2 "
            "/ step-count:3.0, plus a FhirSource row tying Patient 40001 ↔ "
            "Oura. teardown.sh + teardown.sh --purge.",
            "Updated pause_ingest/pause_ingest/exchange.py with the three "
            "real-JHE contract fixes; pause_ingest/pause_ingest/fhir.py "
            "with the derived-features routing.",
            "Updated pause_ingest/examples/oura_sample_upload.py to round-"
            "trip BOTH raw heart-rate (mapped, integer pk) AND derived HRV-"
            "time-domain (auxiliary, UUID pk, with derivedFrom).",
            "pause_ingest/tests/conftest.py — PAUSE_USE_REAL_JHE marker + "
            "collection hook for the mutually-exclusive suites.",
            "pause_ingest/tests/test_exchange_real_jhe.py — 7-test real-"
            "JHE contract suite (passes 7/7 against jhe-local on the ship "
            "date).",
            "docs/JHE_SETUP_RUNBOOK.md, docs/JHE_REAL_RUN_2026-06-16.md "
            "(the transcript), updated pause_ingest/README.md.",
        ],
        "verified": [
            "jhe-local/bootstrap.sh → both containers Up, OAuth token + "
            "FHIR read/write works against the local instance.",
            "pause_ingest/examples/oura_sample_upload.py against jhe-local "
            "round-trips 2 observations (raw heart-rate + derived HRV-time-"
            "domain) and prints OK.",
            "Default pytest: 27/27 (mock suite); PAUSE_USE_REAL_JHE=1 pytest: "
            "7/7 (real_jhe suite). Both modes are mutually exclusive — the "
            "non-active suite is skipped, not failed.",
            "/proposal/integration + /roadmap status pills updated: JHE "
            "designed → prototype.",
        ],
    },
    {
        "title": (
            "Phase 28 — MCP Streamable HTTP transport + MCP Bridge + Agentforce 3.0 Registry"
        ),
        "ask": (
            "Ship the Streamable HTTP transport for the MCP server at "
            "/api/mcp so Agentforce 3.0's MCP Registry can discover and "
            "use Pause's tools. Then turn the Care Router itself into an "
            "MCP host that calls tools from arbitrary external MCP servers."
        ),
        "decisions": [
            "Built the Streamable HTTP transport at /api/mcp using the "
            "MCP SDK's WebStandardStreamableHTTPServerTransport — drops "
            "directly into Next.js App Router's Request → Response shape "
            "with no Express adapter. Stateless mode (sessionIdGenerator "
            "= undefined) since Vercel functions are short-lived and the "
            "Agentforce Registry stores session at its own layer.",
            "Same four tools (timeline / intake / providers / health "
            "check), single registration in frontend/lib/mcp/tools.ts. "
            "The stdio transport in mcp/src/server.ts and the Streamable "
            "HTTP transport in app/api/mcp/route.ts both consume the same "
            "tool registrations — schema-aware, single source of truth.",
            "MCP Bridge (commit f333c2d): turned the Care Router into an "
            "MCP host that loads remote MCP servers per-request, calls "
            "their tools, and surfaces results in the agent-fabric trace. "
            "Loopback remote (Pause's own /api/mcp) is always on so the "
            "Care Router demonstrates the host pattern against our own "
            "tools without a third-party dependency. External remotes are "
            "configurable via PAUSE_MCP_HOST_REMOTES (JSON-encoded array).",
            "Production flip (commit 0dcc65c): MCP Bridge enabled in "
            "production env. Every Care Router decision now records a span "
            "for each MCP tool call it made — fan-out, latency, success/"
            "failure visible in the trace plane.",
        ],
        "built": [
            "frontend/app/api/mcp/route.ts — Streamable HTTP MCP server. "
            "GET/POST/DELETE all handled.",
            "frontend/lib/mcp/tools.ts — shared tool registrations consumed "
            "by both stdio + HTTP transports.",
            "frontend/lib/mcp/host.ts — MCPHost class + resolveRemotesFromEnv "
            "+ createMCPHostFromRequest. Per-request lifecycle (fresh MCP "
            "handshake per Care Router invocation; transport is closed at "
            "end of request).",
            "frontend/lib/mcp/host.test.ts + host.integration.test.ts — 14 "
            "tests pinning the host's fan-out behavior + the loopback path.",
            "Wiring in /api/agents/care-router/tasks to use the host for "
            "tool calls. Span types extended in lib/agent-fabric.ts.",
            "mcp/v0.2.0 published — README + smoke updated for the new "
            "transport.",
        ],
        "verified": [
            "Agentforce 3.0 Registry connection: paste "
            "https://pause-health.ai/api/mcp into Setup → New MCP server → "
            "Pause's four tools enumerate cleanly + are callable from the "
            "Agentforce agent's builder.",
            "Care Router smoke: a routing request makes a loopback MCP "
            "tools/call → the trace shows the span with remoteId: 'loopback' "
            "+ tool name + duration + result attribution.",
            "External-remote test: PAUSE_MCP_HOST_REMOTES set to a "
            "speculative Salesforce MCP endpoint → the host attempts the "
            "call, traces the failure cleanly when the endpoint isn't "
            "reachable, falls back without breaking the routing decision.",
        ],
    },
    {
        "title": (
            "Phase 29 — Headless 360 conformance audit + four gap closures"
        ),
        "ask": (
            "Stand up a Salesforce Headless 360 conformance audit page so "
            "investors can see honestly which of Salesforce's Headless 360 "
            "architectural invariants Pause already satisfies vs which "
            "need work. Then close each gap as a dormant env-gated seam "
            "so the audit moves from 4×designed to 4×prototype."
        ),
        "decisions": [
            "Authored /proposal/headless-360 with three-row pattern-mapping "
            "table (REST + MCP + A2A) showing which Pause surface implements "
            "each Headless 360 pattern, and a four-row gap table naming the "
            "specific invariants the prototype doesn't yet satisfy. Each "
            "gap row has a pill: designed → prototype → shipped.",
            "Gap #1 — PKCE External Client App OAuth flow (commit a20806e). "
            "Authored lib/salesforce-headless360.ts + six routes under "
            "/api/salesforce/headless-360/* (config, authorize, callback, "
            "token/refresh, me, logout). RFC 7636 PKCE with S256, signed-"
            "cookie session envelope with HMAC-SHA256 + timingSafeEqual "
            "tamper detection. Three-state status machine (designed | "
            "prototype | shipped) driven by SF_HEADLESS360_* env vars. 25 "
            "unit tests pin the env matrix, PKCE alphabet, signed-cookie "
            "round-trip, tamper detection.",
            "Gap #2 — `mcp_api` scope on /api/mcp (commit 3202a2a). Built "
            "validateMcpApiBearer with introspect-first (strict: requires "
            "active=true AND scope contains mcp_api) + userinfo fallback "
            "(permissive: verifies token aliveness only when introspect is "
            "disabled — flagged in the return value as via: 'userinfo-"
            "fallback' so callers can log the weaker guarantee). /api/mcp "
            "route handler wraps the validator behind SF_HEADLESS360_"
            "REQUIRE_MCP_AUTH=on. WWW-Authenticate header per RFC 6750. "
            "Loopback bearer propagation in lib/mcp/host.ts attaches the "
            "inbound bearer to the loopback remote only — structural same-"
            "origin guarantee (the loopback URL is built from the request "
            "origin so there's no way to leak cross-origin).",
            "Gap #2 follow-ups (commit ec056ad): bounded process-local "
            "cache for positive introspect results (60s TTL, 1024-entry "
            "LRU-on-insert; only positive results cached so a freshly-"
            "issued token isn't stuck rejected). guardMcpAuth now returns "
            "the validated identity instead of swallowing it; /api/mcp "
            "decorates every successful response with X-Pause-MCP-User + "
            "X-Pause-MCP-Via headers for trace-plane attribution. New "
            "GET /api/mcp/whoami diagnostic endpoint returns "
            "{gate: 'on'|'off', via, username} so operators can verify "
            "gate wiring without parsing the SSE response stream.",
            "Gap #3 — Agent Fabric → Salesforce Platform Event egress "
            "(commit 830bb8e). Initial framing was 'write into Real-Time "
            "Event Monitoring's stream' but research before code corrected "
            "it: RTEM's catalog (LoginEvent, ApiEvent, etc.) is Salesforce-"
            "platform-internal — external apps cannot define new RTEM "
            "event types. The partner-supported pattern is custom Platform "
            "Events. lib/salesforce-platform-event-sink.ts emits each "
            "Agent Fabric span as a Pause_Agent_Trace__e Platform Event "
            "via REST sObjects, authenticated via OAuth 2.0 Client "
            "Credentials against a dedicated Connected App. Fire-and-"
            "forget, never blocks routing, swallows all Salesforce errors. "
            "Token cached for expires_in − 60s; 401 wipes the cache so "
            "the next call re-mints. 22 unit tests + a no-throw-on-failure "
            "invariant test.",
            "Gap #4 — Salesforce CLI parity (commit 0626bf6). Salesforce's "
            "Headless 360 trust model exposes every agent capability "
            "through REST + MCP + CLI. Pause already had REST + MCP; this "
            "completed the triad. Built cli/ in-repo as @pause-health/cli "
            "with four commands (pause health, pause providers, pause "
            "timeline, pause intake) wrapping /api/mulesoft/*. Hand-rolled "
            "argv parser (zero runtime deps) — Salesforce's sf CLI is "
            "itself a heavy tree, but for a four-endpoint wrapper a "
            "minimal parser is honest. --json for jq piping; --base-url + "
            "PAUSE_BASE_URL for preview deploys; PAUSE_API_KEY → "
            "Authorization. 17 unit tests + 6 smoke cases against live "
            "pause-health.ai.",
        ],
        "built": [
            "/proposal/headless-360 — conformance audit page with the "
            "three-row pattern-mapping + four-row gap table + activation "
            "snippets for each gap.",
            "lib/salesforce-headless360.ts — PKCE seam, validateMcpApiBearer "
            "with introspect+userinfo fallback, process-local cache, "
            "isMcpApiAuthRequired, signed-cookie helpers.",
            "frontend/app/api/salesforce/headless-360/* — six routes "
            "implementing the OAuth Authorization Code + PKCE flow.",
            "frontend/app/api/mcp/whoami/route.ts — operator-side gate "
            "diagnostic.",
            "lib/salesforce-platform-event-sink.ts + agent-fabric "
            "integration — Pause_Agent_Trace__e Platform Event egress.",
            "cli/ npm package — @pause-health/cli with four commands, "
            "hand-rolled argv parser, smoke harness against live endpoints.",
            "docs/HEADLESS_360_RUNBOOK.md + docs/SF_PLATFORM_EVENT_SINK_"
            "RUNBOOK.md.",
        ],
        "verified": [
            "/proposal/headless-360 reads: all four gaps prototype (gap #1 "
            "PKCE 2026-06-24, gap #3 Platform Event sink 2026-06-24, gap #2 "
            "mcp_api gate 2026-06-27, gap #4 CLI 2026-06-27).",
            "Gap #2 vitest: 53 tests including the introspect/scope/"
            "inactive/fallback matrix, the cache TTL + negative-no-cache + "
            "per-token-isolation behavior, and the whoami diagnostic envelope.",
            "Gap #4 CLI smoke: 6/6 cases green against live pause-health.ai.",
            "Total project test count after Phase 29: 493 frontend vitest "
            "+ 17 cli vitest = 510 passing (up from 409 before gap #1).",
        ],
    },
    {
        "title": (
            "Phase 30 — MuleSoft Phase 3: nine Anypoint Exchange assets"
        ),
        "ask": (
            "The /proposal/mulesoft page's Phase 3 (multi-customer fabric) "
            "had been pilled `future` since launch — assumed it required a "
            "second customer. That assumption was wrong. Shared System-API "
            "artifacts can land on Exchange the moment they exist, "
            "decoupled from any single customer onboarding."
        ),
        "decisions": [
            "Asset #1 — pause-omh-to-fhir-library v1.0.0 (commit b7f143b). "
            "Promoted the OMH→FHIR R5 Observation DataWeave transform out "
            "of pause-mulesoft-health-v1 into a versioned Anypoint Exchange "
            "library asset. Worker bumped to 1.0.5 and consumes the library "
            "as a Maven dependency at dw::pause::health::omh. Flow XML "
            "unchanged so the runtime response is byte-identical to 1.0.4. "
            "This is the ONLY end-to-end consumed Phase-3 artifact today "
            "(the CloudHub worker actually pulls + bundles it); the rest "
            "are contract-only.",
            "Assets #2-#9 — eight OAS 3.0 spec assets. Two architectural "
            "patterns covered: PULL-FROM-VENDOR (pause-oura-system-api-spec "
            "as the template; pause-whoop and pause-garmin as clones with "
            "vendor-specific data-type catalogs + auth quirks — Whoop has "
            "synthetic OMH recovery-score:1.0 + cardiovascular-strain:1.0 "
            "schemas; Garmin honestly documents OAuth 1.0a upstream + "
            "webhook-pull cadence) and UPLOAD-TO-PAUSE (pause-healthkit-"
            "system-api-spec for iOS-app-side HealthKit batches; pause-"
            "empatica-system-api-spec for researcher-uploaded .zip session "
            "archives; the latter honestly anchored to pause_ingest's "
            "EmpaticaIngestNotImplemented stub blocked by devicely's "
            "numpy<2.0 pin). Plus pause-jhe-system-api-spec (JHE's REST "
            "surface + Django data plane), pause-dbdp-system-api-spec "
            "(HRV feature compute with two modes), and pause-ingest-"
            "process-api-spec (Process tier orchestration; completes the "
            "API-led three-tier story on Exchange).",
            "Two non-obvious gotchas discovered + documented (apply to all "
            "8 spec assets): (1) mvn-deploy-plugin sends .pom uploads with "
            "Content-Type: application/x-www-form-urlencoded by default, "
            "which Exchange v2 500s on with java.io.EOFException: input "
            "contained no data. The .jar upload via mvn works fine (aether "
            "sets octet-stream for jars). Workaround: direct curl PUT for "
            "the POM with Content-Type: application/xml. (2) Tagging a "
            "DataWeave-only library jar with classifier=mule-plugin "
            "triggers Exchange's ms-exchange-tooling-service extension-"
            "model extraction, which 502s on a no-SDK jar. Plain jar "
            "packaging without the classifier is correct — the Mule "
            "runtime picks up the dw/ namespace from any jar on the "
            "classpath.",
            "All eight specs honestly framed as contract-only in their "
            "info.description. Phase 1c will materialize the deployable "
            "Mule projects; pause_ingest does the equivalent orchestration "
            "in-process today. The proposal-page Phase 3 detail names all "
            "nine assets + the two architectural patterns covered.",
        ],
        "built": [
            "mulesoft/pause-omh-to-fhir-library/ — plain Maven jar with "
            "src/main/resources/dw/pause/health/omh.dwl. Published 1.0.0 + "
            "consumed by worker 1.0.5.",
            "mulesoft/specs/pause-{jhe,dbdp,oura,whoop,garmin,healthkit,"
            "empatica,ingest-process}-system-api-spec/ — eight OAS 3.0 spec "
            "assets, each with pom.xml + README + .gitignore + the .yaml "
            "spec file. All eight published to Anypoint Exchange v2 with "
            "status: published.",
            "mulesoft/specs/pause-omh-to-fhir-library/README.md documents "
            "both Exchange-v2 gotchas with the curl-PUT recipe.",
        ],
        "verified": [
            "Every OAS spec validates clean (parses, no missing $refs).",
            "Exchange asset listing for each of the 9 assets returns "
            "status: published.",
            "Worker 1.0.5 deployed to CloudHub Sandbox in 2:12 (BUILD "
            "SUCCESS); direct /health + /providers smoke green; the "
            "deployable mule-application jar bundles "
            "repository/.../pause-omh-to-fhir-library-1.0.0.jar.",
            "Total Phase 3 assets on Exchange: 9 (1 consumed library + "
            "8 contract-only specs).",
        ],
    },
    {
        "title": (
            "Phase 31 — Agentforce Voice partner-web seam"
        ),
        "ask": (
            "Salesforce's Headless 360 product surface includes Agentforce "
            "Voice (the voice channel). Stand up a partner-web seam "
            "matching the same env-driven, dormant-until-activated pattern "
            "the Headless 360 gap closures use."
        ),
        "decisions": [
            "Built the partner-web seam as a separate lib + routes from the "
            "Headless 360 PKCE one, even though both speak OAuth to "
            "Salesforce. Agentforce Voice has its own External Client App "
            "model (different scopes: agentforce_api + voice_call_api), so "
            "fusing them into one would have created a misleading 'one knob "
            "to rule them all' that doesn't actually exist on the Salesforce "
            "side.",
            "Env-driven activation: AGENTFORCE_VOICE_CLIENT_ID + "
            "AGENTFORCE_VOICE_CLIENT_SECRET + AGENTFORCE_VOICE_AUTH_BASE_URL. "
            "Status pill on /proposal/agentforce-voice: designed when "
            "unset, prototype when set + AGENTFORCE_VOICE_VERIFIED is unset, "
            "shipped when both set. Mirrors the Headless 360 three-state "
            "model.",
            "Gated on Salesforce Agentforce Contact Center licensing: the "
            "underlying voice infrastructure requires a SKU Pause doesn't "
            "have on the trailsignup org. Documented the procurement path "
            "in AGENTFORCE_VOICE_RUNBOOK.md for the customer-side "
            "activation. Until that SKU lands, the seam is functionally "
            "designed-only.",
        ],
        "built": [
            "lib/agentforce-voice.ts — config parser + token client + "
            "isAgentforceVoiceConfigured guard + three-state status machine.",
            "frontend/app/api/agentforce-voice/* routes — config probe + "
            "token-mint + a placeholder call endpoint that returns 503 "
            "until Contact Center licensing activates.",
            "/proposal/agentforce-voice — pattern card with the dormant-"
            "seam pill + activation snippet.",
            "docs/AGENTFORCE_VOICE_RUNBOOK.md — procurement steps for the "
            "Contact Center SKU, the External Client App scopes, the "
            "env-var checklist.",
        ],
        "verified": [
            "Config probe with env unset: returns designed.",
            "Config probe with env set + AGENTFORCE_VOICE_VERIFIED unset: "
            "returns prototype + the configured scopes.",
            "Token-mint route 503s with an honest diagnostic when "
            "Contact Center licensing isn't active.",
            "/proposal/agentforce-voice reads as a dormant seam linked "
            "from /proposal/headless-360.",
        ],
    },
    {
        "title": (
            "Phase 32 — Honesty & coverage hardening: drift guards and "
            "route tests across the agent-integration surfaces"
        ),
        "ask": (
            "With the feature surfaces built, spend a run of focused "
            "sessions raising the floor rather than adding features: hunt "
            "down places where an advertised capability, version, or data "
            "claim had drifted from what the code actually does, fix each, "
            "and pin it with a test so it can't drift back; then close the "
            "biggest remaining test-coverage gaps on the highest-stakes "
            "surfaces — the A2A handoff, the Agent Fabric, the MCP tool "
            "plane, the MuleSoft Experience APIs, and the Salesforce "
            "Headless 360 OAuth seam. Ground rule: test-only unless a fix "
            "is required, and every drift fix ships with a guard."
        ),
        "decisions": [
            "Adopted a 'guard, don't couple' pattern for cross-surface "
            "consistency. Where two representations of the same fact live "
            "in different modules (a registry entry vs the real server "
            "version; a public descriptor vs the registered tool names; a "
            "hardcoded dataset count vs the generated meta file), the fix "
            "was NOT to import one into the other — that would drag heavy "
            "dependencies across package boundaries. Instead a parity test "
            "asserts they match, so a future edit to one without the other "
            "fails CI. This mirrors the existing tools.parity approach.",
            "Made the Agent Fabric registry derive each agent's policy list "
            "from the policy catalog (POLICIES[].appliesTo) as a single "
            "source of truth, rather than a hand-maintained array that had "
            "drifted — the Care Router advertised policies it didn't "
            "enforce AND named one that didn't exist. The public A2A Agent "
            "Card now derives the same way, so the discovery document can't "
            "overclaim.",
            "Fixed three concrete drifts and guarded each: the Agent Fabric "
            "registry advertised a stale Pause MCP version (0.1.0 vs the "
            "server's real 0.3.0); the PUBLIC .well-known/mcp.json — what "
            "external MCP clients read — had re-drifted to the same stale "
            "0.1.0 AND claimed providers are 'ranked by Pause's internal "
            "graph score' when the code ranks distance-first; and the MCP "
            "provider tool's LLM-facing description hardcodes dataset counts "
            "(2,015 providers / 1,720 sanctioned-filtered) that nothing tied "
            "to the generated meta.",
            "Made the inbound A2A /tasks handler tolerant of the current "
            "Google A2A spec: the Part discriminator was renamed from `type` "
            "to `kind`, so a spec-current external client tagging its intake "
            "part kind:'data' was being silently ignored (intake collapsed "
            "to {} and the task was rejected by the red-flag policy for the "
            "wrong reason). Added defensive readers that accept either "
            "discriminator while Pause keeps EMITTING the older `type` form "
            "for its own agents; the empty-intake fail-closed property is "
            "preserved.",
            "Fixed a smoke-test false green: the flagship end-to-end checks "
            "judged A2A calls on HTTP 200 + JSON-parse only, but the A2A "
            "layer returns 200 for governance BLOCKS too, so a blocked "
            "routing looked identical to a completed one. Two payload bugs "
            "meant the checks were hitting the block path every run. Added a "
            "per-call validate() hook so a 200 is necessary but no longer "
            "sufficient — the multi-agent cases now assert the task actually "
            "completed with a RoutingDecision.",
            "Chose an in-process integration rig for the MCP tool + route "
            "tests: a real McpServer connected to a real Client over an "
            "InMemoryTransport, with an injected recording fetch. Every "
            "assertion rides a genuine tools/call round-trip rather than a "
            "re-implementation, and the /api/mcp Streamable-HTTP route is "
            "driven with hand-built JSON-RPC frames whose SSE responses are "
            "parsed — which surfaced that the tool's async fetch only fires "
            "as the SSE body is consumed (the same reason the route must not "
            "eagerly close the transport).",
            "Treated the Salesforce Headless 360 OAuth routes as the "
            "highest-value coverage gap: the library primitives (PKCE, "
            "cookie signing, validateMcpApiBearer) were deeply tested, but "
            "the six routes wiring them into the actual flow — including the "
            "PKCE authorize kickoff and the CSRF-guarded callback — had zero "
            "tests. Prioritized the security-load-bearing invariants: state "
            "binding, cookie tamper-evidence, hardened cookie flags, and "
            "open-redirect defense.",
        ],
        "built": [
            "lib/a2a.ts — partKind() + findDataPart() spec-tolerant Part "
            "readers; lib/a2a.test.ts covering the client, helpers, and "
            "both discriminators; /tasks route now extracts intake via "
            "findDataPart with a route test proving a kind:'data' task "
            "completes.",
            "Agent Fabric single-sourcing: withPolicies() derives each "
            "registry entry's policies from the catalog; evaluateGovernance "
            "enforces the rationale-required policy; the Care Router Agent "
            "Card derives its advertised policies from getPoliciesForAgent. "
            "agent-fabric.test.ts + the four agent-fabric API route tests "
            "(governance/evaluate, policies, traces, sf-sink/config) + the "
            "Agent Card contract test.",
            "MCP drift guards: registry-parity.test.ts (Agent Fabric entry "
            "⇄ SERVER_VERSION + registered tool names), "
            "public-descriptor-parity.test.ts (.well-known/mcp.json ⇄ same, "
            "plus a ranking-honesty assertion), and "
            "tools-provider-counts.parity.test.ts (the description's numbers "
            "⇄ provider-directory.generated.meta.json). Corrected the "
            "registry version, the public descriptor version + ranking "
            "claim.",
            "MCP behavior + wiring: tools.behavior.test.ts (all four tool "
            "handlers — URL construction, headers, summary narration, error "
            "path, over a real in-memory transport); host.test.ts coverage "
            "of createMCPHostFromRequest header parsing + the cross-origin "
            "bearer-leak guard; http-auth.test.ts pinning guardMcpAuth's "
            "discriminated result + attachIdentityHeaders; and "
            "app/api/mcp/route.test.ts (Streamable-HTTP initialize / "
            "tools/list / tools/call round-trip, request-origin base-URL "
            "derivation, and all three verbs running through the auth gate).",
            "MuleSoft Experience API route tests: the previously-untested "
            "patient/[id]/timeline + patient/[id]/intake routes (FHIR bundle "
            "/ intake shape, meta bookkeeping, id-aliasing, cache headers).",
            "Salesforce Headless 360 route tests for all six OAuth endpoints "
            "— authorize (PKCE params, hardened cookie flags, state⇄verifier "
            "binding, open-redirect defense), callback (CSRF state-mismatch "
            "refusal, tamper defense, token exchange, session cookie), me "
            "(session-state ladder), token/refresh (cookie-carried token, "
            "clear-on-failure), config (no-secret-leak probe), and logout.",
            "scripts/smoke-test.mjs — shared well-formed SMOKE_INTAKE + a "
            "per-call validate() hook so the A2A + handoff cases assert real "
            "completion, not merely a 200.",
        ],
        "verified": [
            "Frontend Vitest suite grew from the low-500s to 654 tests "
            "across 53 files over the pass, all green; next build clean "
            "after every increment.",
            "Each drift guard was confirmed to actually bite via a negative "
            "control — e.g. bumping the descriptor version or the generated "
            "provider total makes the corresponding parity test fail before "
            "restoring.",
            "The kind:'data' A2A tolerance was verified live against a fresh "
            "production build: kind:'data' now returns state:'completed' / "
            "decision:'allow' with a RoutingDecision, type:'data' is "
            "unchanged, and a text-only message still fails closed.",
            "The security-critical Headless 360 invariants are pinned: a "
            "mismatched OAuth state is refused 400 and the token endpoint is "
            "never contacted; a tampered pending/session cookie fails "
            "verification; refreshed tokens are written to the cookie and "
            "never returned in the response body.",
            "Every change in the pass was test-only except the four targeted "
            "drift corrections (registry version, descriptor version + "
            "ranking, A2A part tolerance) and the smoke-test fix; all shipped "
            "with their guards and a changelog entry.",
        ],
    },
]

OPERATIONS_LOG = {
    "title": "Operations log — recurring patterns across the build",
    "items": [
        {
            "name": "Dev server stale-cache failures (Cannot find module './<id>.js')",
            "detail": (
                "Repeatedly observed during long sessions: Next.js dev "
                "occasionally drifts its .next webpack runtime out of sync "
                "with the on-disk routes after large structural changes. "
                "Resolution that worked every time: kill the dev server "
                "(lsof -ti:3000 | xargs kill -9), rm -rf frontend/.next, "
                "restart with WATCHPACK_POLLING=true CHOKIDAR_USEPOLLING=1 "
                "after raising ulimit -n. Documented in README."
            ),
        },
        {
            "name": "EMFILE / file-descriptor limits on macOS",
            "detail": (
                "The Next.js watcher plus dev-mode hot-reload plus polling "
                "exceeds the default macOS open-files cap. ulimit -n 8192 "
                "(or 10240) per shell session before npm run dev resolves "
                "it. Polling watchers (WATCHPACK_POLLING + CHOKIDAR_USE"
                "POLLING) used as a belt-and-suspenders measure."
            ),
        },
        {
            "name": "Module-scoped global for cross-route state",
            "detail": (
                "Both lib/agent-fabric.ts and the Data 360 mock state need "
                "to survive Next dev hot reload AND be the same instance "
                "across every API route. Pattern: store the state object "
                "on globalThis under a private key, lazily initialize, "
                "type the global with a one-line type extension. Lets one "
                "API route record a span that a different API route can "
                "read in the same process."
            ),
        },
        {
            "name": "Soft / dynamic imports for optional SDKs",
            "detail": (
                "Anthropic SDK is imported via `await import('@anthropic"
                "-ai/sdk')` inside the Claude code path so the rest of the "
                "build does not depend on the SDK being installed. The "
                "scripted-policy fallback runs even if @anthropic-ai/sdk "
                "is missing. Same pattern works for any future agent SDK."
            ),
        },
        {
            "name": "Environment-variable gating everywhere",
            "detail": (
                "Every external integration follows the same pattern: read "
                "an env var; if present, run the real path; if absent, run "
                "a deterministic fallback that produces a payload of the "
                "same shape. Documented in .env.example. Lets a reviewer "
                "run the entire prototype with zero credentials."
            ),
        },
        {
            "name": "npm install inside nested monorepo packages",
            "detail": (
                "Initial mcp/ install failed because npm walked up to find "
                "the repo's root and there is no root package.json (this "
                "is intentional). Resolution: cd mcp && npm install "
                "--prefix . explicitly scoped to the package directory."
            ),
        },
        {
            "name": "Server-to-server OAuth Client Credentials as the default integration shape",
            "detail": (
                "Salesforce Health Cloud (Phase 12), Salesforce Data "
                "Cloud (Phase 14, scope-add only), and the recommended "
                "MuleSoft Anypoint integration (Phase 16) all use the "
                "same auth shape: server-side OAuth 2.0 Client "
                "Credentials Flow via a per-integration Connected/External "
                "Client App, cached token in process memory, in-flight "
                "request deduplication. lib/salesforce/auth.ts is the "
                "canonical template; the same ~120 lines will work for "
                "Anypoint by swapping the token URL and scope names. "
                "Avoid embedding-style integrations (browser widgets, "
                "iframe-based UIs) — those drag in CORS, CSP "
                "frame-ancestors, and Experience-site restrictions that "
                "cost a session each."
            ),
        },
        {
            "name": "Zscaler intercepts on Salesforce-edge hostnames",
            "detail": (
                "Corporate Zscaler proxy intercepted *.c360a.salesforce"
                ".com (Data Cloud) with a 504 Gateway Timeout during "
                "Phase 14 investigation. Pausing Zscaler resolved it. "
                "Anypoint Platform hostnames resolve through the SAME "
                "*.edge2.salesforce.com edge family (verified in Phase "
                "16) so the same Zscaler posture covers both. Risk "
                "documented: if Zscaler tightens *.salesforce.com later, "
                "Salesforce AND MuleSoft integrations break together — "
                "treat them as one dependency for security-posture "
                "conversations."
            ),
        },
        {
            "name": "Embedded widgets require deployment in the embedding org",
            "detail": (
                "Phase 15 invested a full session investigating why the "
                "SDO sample Agentforce deployment could not be embedded "
                "on pause-health.ai. Phase 18 confirmed the underlying "
                "constraint: embedded-widget integrations need a "
                "deployment authored in an org whose Experience site can "
                "be configured to include the embedding origin in BOTH "
                "(a) the SCRT2 endpoint's CORS allowlist (driven by "
                "Trusted URLs + CorsWhitelistEntry on the embedding org) "
                "AND (b) the site's frame-ancestors CSP (auto-derived "
                "from the Experience site's Trusted Domains, which is "
                "in turn populated from the Domain field set during "
                "EmbeddedServiceDeployment creation). The SDO sample "
                "had neither knob configured for pause-health.ai, and "
                "neither could be changed from outside the org. A "
                "Pause-Health-owned deployment in the same org (Phase "
                "18) trivially has both knobs because we control the "
                "Domain field during creation. Lesson: don't try to "
                "borrow a sample deployment for external embedding; "
                "stand up your own deployment, takes ~30 min once you "
                "know the path."
            ),
        },
        {
            "name": "EmbeddedServiceConfig.clientVersion defaults to WebV1",
            "detail": (
                "Phase 18 surfaced a non-obvious gotcha: even when "
                "DeploymentFeature is set to EmbeddedMessaging (the V2 "
                "Messaging for In-App and Web product), the underlying "
                "EmbeddedServiceConfig.Metadata.clientVersion field "
                "defaults to 'WebV1'. The SCRT2 runtime config endpoint "
                "reads this field verbatim and the SDK switches between "
                "V1 and V2 wire protocols based on its value. With "
                "clientVersion=WebV1 the launcher renders but the chat "
                "panel handshake fails with RPC connection timeout. "
                "Fix: PATCH Metadata.clientVersion='WebV2' via Tooling "
                "API v65 (older API versions do not expose the field at "
                "all), then republish the deployment. Document in any "
                "future runbook as a mandatory step."
            ),
        },
        {
            "name": "Agentforce Service Agents do not appear in Einstein Bots",
            "detail": (
                "Phase 18 surfaced this footgun: the legacy 'Einstein "
                "Bots' page in Salesforce Setup lists only classic "
                "bots, not Agentforce Service Agents. An Agentforce "
                "agent created via Agent Builder exists in BotDefinition "
                "(queryable via SOQL) but is invisible in the Einstein "
                "Bots UI. To manage Agentforce agents go to Setup -> "
                "Agentforce Agents (or the agent's detail page directly). "
                "Anyone porting from older Salesforce docs will look at "
                "the empty Einstein Bots page and assume their agent "
                "wasn't created — this is wrong."
            ),
        },
        {
            "name": "Bind Bot-to-Channel on the Channel, not on the Agent",
            "detail": (
                "The Agentforce agent's detail page has a 'Connections' "
                "tab. It looks like the place to bind the agent to a "
                "messaging channel. It is NOT — that tab only supports "
                "Type=API (external app integrations). The actual "
                "Bot-to-MessagingChannel binding lives on the channel: "
                "Setup -> Messaging Settings -> <channel> -> "
                "Omni-Channel Routing -> Routing Type = 'Agentforce "
                "Service Agent' -> pick the agent. Without this, "
                "incoming conversations route to the channel's "
                "FallbackQueue (typically an empty 'Messaging' queue) "
                "and the chat panel sits forever at RPC connection "
                "timeout."
            ),
        },
        {
            "name": "Investigate-only sessions with a runbook deliverable",
            "detail": (
                "Phases 14, 15, and 16 each ended with a deferred "
                "implementation and a deliverable runbook (PHASE_3_"
                "RUNBOOK.md, MULESOFT_RUNBOOK.md, or in-line notes in "
                ".env.example/.env.local for Phase 14). The pattern: "
                "ask up front for scope (investigate-only vs investigate"
                "+execute vs full session), do the discovery, write the "
                "runbook so the NEXT session is click-through only, "
                "commit the runbook (docs are cheap to roll back). "
                "Avoids the failure mode of starting an implementation, "
                "hitting a hard ceiling mid-session, and leaving the "
                "next agent or developer to rediscover the same "
                "blockers."
            ),
        },
        {
            "name": "Commit discipline",
            "detail": (
                "One commit per coherent feature. Every commit message "
                "follows the same shape: title line, blank line, motivation, "
                "blank line, bulleted list of new surface and wiring. Commit "
                "history reads as a build journal on its own."
            ),
        },
        {
            "name": "Anypoint Exchange v2 POM Content-Type 500",
            "detail": (
                "Discovered in Phase 30. mvn-deploy-plugin sends .pom "
                "uploads with Content-Type: application/x-www-form-"
                "urlencoded by default; Anypoint Exchange v2's Maven "
                "endpoint responds with 500 + 'java.io.EOFException: "
                "input contained no data'. The .jar upload via mvn works "
                "fine (aether sends application/java-archive for jars). "
                "Workaround that ships: direct curl PUT for the POM with "
                "Content-Type: application/xml. Recipe documented in "
                "every Phase-3 spec asset's README.md. Affects ALL 8 "
                "spec assets + the DataWeave library."
            ),
        },
        {
            "name": "Don't tag DataWeave-only libraries classifier=mule-plugin",
            "detail": (
                "Discovered in Phase 30. Anypoint Exchange's "
                "ms-exchange-tooling-service tries to extract a Mule "
                "extension model from any artifact tagged "
                "classifier=mule-plugin; a no-SDK jar (one containing "
                "only DataWeave resources) makes it 502 with "
                "'BadGatewayError: invalid json response body: Error "
                "proc...'. The Mule runtime picks up the dw/ namespace "
                "from any jar on the classpath, so plain jar packaging "
                "without the classifier is correct. Use "
                "classifier=mule-plugin only for real Mule SDK extensions "
                "(Custom Connectors built with Studio's packager)."
            ),
        },
        {
            "name": "vi.fn().mockResolvedValue(Response) and consumed bodies",
            "detail": (
                "Caught in Phase 29 gap #2 follow-ups. "
                "vi.fn().mockResolvedValue(new Response(...)) returns the "
                "SAME Response object on every call — but a Response body "
                "can only be consumed once with .json(). If a test calls "
                "the validator twice with mockResolvedValue, the second "
                ".json() throws and the validator falls through to a "
                "different code path (in our case, the introspect → "
                "userinfo fallback), making fetch call counts off by "
                "one. Fix: vi.fn().mockImplementation(async () => new "
                "Response(...)) for any test that drives the validator "
                "more than once. Pre-existing tests didn't hit this "
                "because each called validateMcpApiBearer once; the new "
                "TTL + negative-cache tests called it twice and revealed "
                "the issue. Documented in memory/project_headless360_"
                "state.md."
            ),
        },
        {
            "name": "Free-tier ngrok allows one tunnel; CF Tunnel needs a domain",
            "detail": (
                "Surfaced when scoping Phase 1c. Pause's Flex Gateway "
                "(Phase 23 iteration 3) already owns the "
                "cattail-reactive-sassy.ngrok-free.dev pinned subdomain; "
                "a second ngrok process can't get the same domain on "
                "free tier. Cloudflare Tunnel sounds like a drop-in "
                "alternative but the 'free permanent domain' story only "
                "applies if you already own a domain on CF — "
                "trycloudflare.com URLs are ephemeral and change per "
                "restart. The architecturally correct next move is the "
                "Phase-1 iteration-8 persistent VM (Fly.io / Lightsail) "
                "hosting the Flex Gateway, which also unblocks Phase 1c. "
                "Documented as the 'iteration 8' row in docs/MULESOFT_"
                "RUNBOOK.md."
            ),
        },
    ],
}

CURRENT_STATE = {
    "title": "Current state — what is live, what is mocked",
    "live": [
        "Marketing site, About page (with founder picture, LinkedIn, "
        "and Person JSON-LD), blog and press scaffolds, footer, legal "
        "pages, SEO surface (sitemap, robots, OG, Twitter, security.txt, "
        "Organization JSON-LD), canonical URLs on the pause-health.ai "
        "apex domain.",
        "Investor brief: 16 deep-dive pages plus the /proposal/full "
        "single-document narrative — including /proposal/headless-360 "
        "(four-row conformance audit with all four gaps at prototype) "
        "and /proposal/agentforce-voice (Headless 360 voice surface).",
        "Clickable prototype: /demo/intake (Agentforce real-or-scripted "
        "fallback with persona picker + visible Pre-Brief Panel), "
        "/demo/patient (persona-aware care detail), /demo/routing "
        "(with live Care Router decision card + distance-ranked "
        "recommended providers), /demo/analytics (persona-filterable "
        "operational metrics + pathway chart), /demo/agent-fabric "
        "(persona-filterable trace console with real-vs-mock source "
        "banners on every span). All five share the DemoShell + "
        "PersonaJourneyFooter and preserve personaId across pages.",
        "Real Code Repository nav link to the GitHub repo (Apache-2.0).",
        "Python wearable-ingest worker (pause_ingest/) — FLIRT-based, "
        "pytest-covered, with the PAUSE_USE_REAL_JHE=1 opt-in suite that "
        "passes 7/7 against the local jhe-local Docker stack.",
        "JupyterHealth Exchange local stack (jhe-local/) — idempotent "
        "Docker-compose bootstrap; OAuth client + Study + Patient + "
        "DataSource + Scope + FhirSource all seeded so pause_ingest can "
        "round-trip both raw + derived FHIR Observations against real JHE.",
        "MCP servers (two transports, one tool registration): mcp/ ships "
        "@pause-health/mcp v0.3.0 over stdio (Claude Desktop, Cursor); "
        "frontend/app/api/mcp serves the same four tools over Streamable "
        "HTTP for the Agentforce 3.0 Registry. Both consume "
        "frontend/lib/mcp/tools.ts so a tool change ships to both at once.",
        "MCP Bridge — the Care Router is itself an MCP host. Loads "
        "remote MCP servers per-request, calls their tools, records "
        "each call as a span. Loopback (Pause's own /api/mcp) always on; "
        "external remotes via PAUSE_MCP_HOST_REMOTES.",
        "Pause CLI (cli/) — @pause-health/cli with four commands "
        "wrapping /api/mulesoft/* (health, providers, timeline, intake). "
        "Hand-rolled argv parser (zero runtime deps), --json mode for "
        "jq piping, --base-url / PAUSE_BASE_URL for preview deploys. "
        "Closes Headless 360 audit gap #4 (REST + MCP + CLI triad).",
        "Multi-agent control plane: real A2A handoff, real Agent Fabric "
        "trace store, real governance evaluation, real Anthropic path "
        "when configured. Trace spans egress to Salesforce as "
        "Pause_Agent_Trace__e Platform Events via the dormant Headless "
        "360 gap-#3 sink when SF_PLATFORM_EVENT_* env vars are set.",
        "Salesforce Health Cloud grounding (Phase 12) — LIVE against a "
        "real connected Developer Edition org via OAuth client "
        "credentials. SOQL against Contact + CareProgramEnrollee + "
        "CarePlan + Case. Six seeded demo Contacts with menopause-"
        "specific care plans. Agent Fabric trace spans show _source: "
        "'real' on the federated-query span when env vars are set; "
        "fall back to mock when unset (zero-credential default).",
        "Salesforce Agentforce Embedded Messaging intake on /demo/intake. "
        "Pause_Health_Intake_Agent (real Agentforce Service Agent on "
        "Service Cloud) responds to messages from the chat panel "
        "embedded via the V2 Messaging-for-Web bootstrap. Routing: "
        "Omni-Channel → Pause_Intake_Prechat_Router Flow → "
        "Pause_Health_Intake_Agent. Phase 18a/b laid the 5-component "
        "data pipeline (channel customParameters → routing Flow → "
        "MessagingSession.Pause_*__c → Bot contextVariables → topic "
        "instructions); Phase 18c discovered the V2 SDK's prechatAPI "
        "was an empty Proxy and pivoted to a visible Pre-Brief Panel; "
        "Phase 26's auto-ZIP work confirmed the V2 prechatAPI is now "
        "FIXED on the live SDK — setHiddenPrechatFields({Patient_Zip}) "
        "actually transports the value through the channel + Flow to "
        "MessagingSession.Pause_Patient_Zip__c, which the agent reads "
        "as $Context.Pause_Patient_Zip and hard-binds into the "
        "findMenopauseProviders action so the agent never re-asks for "
        "ZIP. Agentforce provider-lookup action (Phase 26) lets the "
        "agent call /api/mulesoft/providers via an External Service + "
        "Named Credential. Scripted Pause-branded fallback runs when "
        "the four NEXT_PUBLIC_AGENTFORCE_* env vars are unset.",
        "Salesforce Data Cloud Phase 2 — three Calculated Insights "
        "(HRV, vasomotor burden, sleep disruption) live on trailsignup "
        "org over ssot__Individual__dlm. /api/data-360/patient/<id>/"
        "grounding returns 'Phase 2: SOQL (Health Cloud) + Data Cloud "
        "Calculated Insights' on every call. Phase 2-bis (real DBDP "
        "feature math via the Ingestion API push path) is SHIPPED in "
        "the repo; operationalization for trailsignup is the remaining "
        "Setup UI work in PHASE_2_INGESTION_API_RUNBOOK.md.",
        "MuleSoft Anypoint Phase 1 LIVE on CloudHub 2.0 — worker "
        "pause-mulesoft-health-v1 v1.0.5 serving /health + /providers "
        "behind Flex Gateway (Auth0-JWT validation + plain rate "
        "limiting), DataWeave OMH→FHIR on /health, Phase-2 full "
        "provider contract on /providers (distance-ranked, plan-"
        "narrowed, sanctioned-filtered, MSCP-overlay-flagged). "
        "Production /api/mulesoft/{health,providers} both report "
        "meta._source: 'live-mulesoft' end-to-end. Graceful "
        "degradation to mock-fallback on tunnel-down (laptop sleep, "
        "VPN, etc.); the proxy + worker stay byte-shape-compatible.",
        "MuleSoft Phase 3 — nine Anypoint Exchange assets on the "
        "Pause Health business group. One CONSUMED LIVE: "
        "pause-omh-to-fhir-library v1.0.0 (CloudHub worker 1.0.5 "
        "pulls + bundles it). Eight CONTRACT-ONLY: pause-jhe-system-"
        "api-spec, pause-dbdp-system-api-spec, pause-oura/whoop/"
        "garmin/healthkit/empatica-system-api-spec (covering both "
        "pull-from-vendor + upload-to-Pause patterns), pause-ingest-"
        "process-api-spec. With the Process tier published, the "
        "MuleSoft API-led three-tier story (System + Process + "
        "Experience) is fully on Exchange.",
        "Provider graph LIVE — 2,015 NPPES-derived providers across "
        "all 50 states + DC + 930 ZIP-3 prefixes, 15 MSCP-certified + "
        "2,000 menopause-relevant non-certified. Six NPPES service-"
        "line signals (FACOG/FAAFP/WHNP/multi-taxonomy/etc.). Three "
        "state license-sanction overlays at build time (588 CA + 849 "
        "NY + 283 TX = 1,720 sanctioned candidates dropped before any "
        "ranking). Census 2020 ZCTA centroid distance ranking. "
        "Real-shaped synthetic insurance acceptance per NPI (8 "
        "canonical plan tokens). Browseable /provider directory + "
        "/provider/[npi] profile pages. The Care Router consumes the "
        "SAME queryProviderDirectory function the /provider UI calls "
        "— triage and the directory stay in lockstep.",
        "Headless 360 audit (Phase 29) — /proposal/headless-360 with "
        "the three-row pattern-mapping (REST + MCP + A2A) and the "
        "four-row gap table. ALL FOUR GAPS at prototype: gap #1 PKCE "
        "External Client App seam + 6 routes, gap #2 mcp_api bearer "
        "gate on /api/mcp with introspect+userinfo fallback + 60s "
        "cache + identity-headers + /api/mcp/whoami diagnostic, gap "
        "#3 Platform Event egress sink, gap #4 @pause-health/cli. "
        "Activation per gap is operator-side env-var procurement.",
        "Test coverage: 493 frontend vitest + 17 cli vitest + 27 (mock) "
        "or 7 (real_jhe opt-in) pause_ingest pytest = 510-537 passing "
        "depending on the JHE-suite mode. ESLint + tsc clean across "
        "frontend + cli. Smoke harness: 168/168 routes 200 OK on "
        "production.",
        "GitHub Actions: frontend-check, codeql, dependabot, vercel-"
        "preview, lighthouse-nightly (with checked-in summary.json "
        "time series).",
        "Production deployment on Vercel at pause-health.ai. Open-source "
        "(Apache-2.0) with CONTRIBUTING.md, CODE_OF_CONDUCT.md, "
        "SECURITY.md, NOTICE.",
    ],
    "mocked": [
        "Salesforce Data Cloud Phase 2-bis (real wearable feature math). "
        "Phase 2 SHIPPED the three Calculated Insights (HRV, vasomotor, "
        "sleep) over ssot__Individual__dlm in MAX(constant) form; Phase "
        "2-bis upgraded the CI SQL to actual SUM/AVG over per-patient "
        "rows pushed via the Data Cloud Ingestion API. Push client + "
        "DLO schema + real CI SQL are all SHIPPED in the repo; "
        "operationalizing it for the trailsignup org is the remaining "
        "Setup UI work documented in PHASE_2_INGESTION_API_RUNBOOK.md.",
        "MuleSoft Experience APIs beyond /health + /providers — "
        "/timeline + /intake routes remain Next.js mocks. /health + "
        "/providers are LIVE on CloudHub 2.0 (iterations 1-7 SHIPPED, "
        "with worker v1.0.5 consuming pause-omh-to-fhir-library v1.0.0 "
        "as a Maven dependency, behind Flex Gateway with JWT auth + "
        "rate limiting). The other two endpoints follow the same "
        "live-or-mock proxy pattern; flip with one env var when their "
        "Mule flows ship.",
        "Phase 1c implementations of the Phase-3 spec assets. Eight "
        "OAS 3.0 contracts SHIPPED on Anypoint Exchange (JHE, DBDP, "
        "Oura, Whoop, Garmin, HealthKit, Empatica E4 system APIs + the "
        "ingest-process API), all CONTRACT-ONLY. The deployable Mule "
        "projects that honor each contract are Phase 1c; pause_ingest's "
        "Python worker does the equivalent orchestration in-process "
        "today.",
        "MuleSoft Agent Fabric — in-memory registry, policy catalog, "
        "trace store. Production swaps for the real Anypoint Agent "
        "Fabric service. Spans now egress to Salesforce as "
        "Pause_Agent_Trace__e Platform Events (Headless 360 audit "
        "gap #3) when SF_PLATFORM_EVENT_* env vars are set; without "
        "those env vars the spans live in-process only.",
        "Headless 360 conformance gaps activation. All 4 gaps SHIPPED "
        "as dormant env-gated seams (pill: prototype). Activation per "
        "gap is operator-side env-var procurement: SF_HEADLESS360_* "
        "for gap #1 (PKCE) + gap #2 (mcp_api bearer gate); "
        "SF_PLATFORM_EVENT_* for gap #3 (Platform Event egress); the "
        "@pause-health/cli npm-publish decision for gap #4 (the artifact "
        "ships in-repo today, npm-scope ownership is a separate ops step).",
        "Anthropic Claude — real when ANTHROPIC_API_KEY is set, "
        "deterministic policy engine fallback otherwise.",
        "Agentforce Voice (Phase 31) — partner-web seam SHIPPED dormant; "
        "voice round-trip gated on Salesforce Agentforce Contact Center "
        "licensing not present on the trailsignup org.",
    ],
    "deferred_with_runbook": [
        "Phase 1c — turn one Phase-3 spec asset (pause-jhe-system-api-"
        "spec is the recommended first target) into a deployable Mule "
        "project that wraps pause_ingest.exchange over HTTP. Real "
        "infrastructure constraint: free-tier ngrok allows one tunnel "
        "at a time, and the Flex Gateway already owns the "
        "cattail-reactive-sassy subdomain. Cloudflare Tunnel free tier "
        "needs a CF-owned domain for permanent URLs. The architecturally "
        "right next move is Phase-1 iteration 8 — host the Flex Gateway "
        "on a small persistent VM (Fly.io / Lightsail, ~$5/mo) so it "
        "stops dropping on laptop sleep AND it unblocks the Phase 1c "
        "tunnel constraint. Documented as the 'iteration 8' row in "
        "docs/MULESOFT_RUNBOOK.md.",
        "Headless 360 gap #2 known limitation — userinfo-fallback path "
        "doesn't enforce scope. Operators who need RFC 7662-strict "
        "scope validation must keep introspect enabled on their "
        "External Client App. Salesforce-side org-config choice; "
        "Pause can't close from this side.",
        "Phase 31 — Agentforce Voice voice round-trip. Seam shipped "
        "dormant; voice infrastructure requires Salesforce Agentforce "
        "Contact Center licensing on the customer org. Documented in "
        "docs/AGENTFORCE_VOICE_RUNBOOK.md.",
        "Smoke harness against the live MuleSoft Phase 1 worker — the "
        "frontend smoke at frontend/scripts/smoke-test.mjs covers the "
        "Next.js proxy layer (live-or-mock); a complementary harness "
        "that probes the gateway + worker independently would catch "
        "drift the proxy-layer smoke can mask. ~1 hour.",
    ],
}


# ---------- styling helpers ----------------------------------------------

BRAND_COLOR = RGBColor(0x6A, 0x1B, 0x9A)  # Pause purple
HEADING_COLOR = RGBColor(0x2D, 0x2D, 0x2D)
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)


def _set_paragraph_spacing(p, space_before: float = 4, space_after: float = 6):
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)


def _add_run(p, text: str, *, bold=False, italic=False, size: float = 11,
             color: RGBColor | None = None, monospace: bool = False):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if monospace:
        font.name = "Menlo"
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Menlo")
        rFonts.set(qn("w:hAnsi"), "Menlo")
        rFonts.set(qn("w:cs"), "Menlo")
        rPr.append(rFonts)
    return run


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, text, bold=True, size=26, color=BRAND_COLOR)
    _set_paragraph_spacing(p, 0, 8)


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_run(p, text, italic=True, size=12, color=MUTED_COLOR)
    _set_paragraph_spacing(p, 0, 14)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_run(p, text, bold=True, size=18, color=HEADING_COLOR)
    _set_paragraph_spacing(p, 18, 8)


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_run(p, text, bold=True, size=14, color=HEADING_COLOR)
    _set_paragraph_spacing(p, 12, 4)


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_run(p, text, bold=True, size=12, color=HEADING_COLOR)
    _set_paragraph_spacing(p, 10, 2)


def add_paragraph(doc: Document, text: str, *, italic: bool = False,
                  color: RGBColor | None = None):
    p = doc.add_paragraph()
    _add_run(p, text, size=11, italic=italic, color=color)
    _set_paragraph_spacing(p, 2, 6)


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_run(p, item, size=11)
        _set_paragraph_spacing(p, 1, 2)


def add_monospace_block(doc: Document, text: str):
    """Render preformatted text in a monospace font, no wrapping changes."""
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        _add_run(p, line if line else "\u00a0", size=9, monospace=True)
        _set_paragraph_spacing(p, 0, 0)


def add_page_break(doc: Document):
    doc.add_page_break()


# ---------- rendering ----------------------------------------------------

def render(doc: Document):
    # cover
    add_title(doc, DOC_TITLE)
    add_subtitle(doc, DOC_SUBTITLE)
    add_paragraph(doc, DOC_AUTHOR_LINE, italic=True, color=MUTED_COLOR)

    add_h2(doc, "How to read this document")
    add_paragraph(
        doc,
        "The document is organized as a series of build phases in the "
        "order they happened. Each phase records the ask, the decisions "
        "and trade-offs made, what was built, and how it was verified. "
        "Two appendices at the end summarize recurring operational "
        "patterns and the current state of the prototype.",
    )

    add_h2(doc, "Project overview")
    for para in OVERVIEW_PARAGRAPHS:
        add_paragraph(doc, para)

    add_page_break(doc)

    add_h1(doc, "End-state architecture")
    add_paragraph(
        doc,
        "The diagram below shows the current shape of the prototype after "
        "the most recent (Data 360) phase. Every block on the diagram is "
        "either a real running surface in the repo or a documented mock "
        "with the same wire shape as its production replacement.",
    )
    add_monospace_block(doc, ARCH_DIAGRAM)

    add_h2(doc, "Monorepo layout")
    add_monospace_block(doc, MONOREPO_LAYOUT)

    add_page_break(doc)

    add_h1(doc, "Build phases")

    for i, phase in enumerate(PHASES, start=1):
        add_h2(doc, phase["title"])

        add_h3(doc, "Ask")
        add_paragraph(doc, phase["ask"])

        add_h3(doc, "Decisions and trade-offs")
        add_bullets(doc, phase["decisions"])

        add_h3(doc, "What was built")
        add_bullets(doc, phase["built"])

        add_h3(doc, "Verification")
        add_bullets(doc, phase["verified"])

        if i != len(PHASES):
            add_page_break(doc)

    add_page_break(doc)

    add_h1(doc, OPERATIONS_LOG["title"])
    add_paragraph(
        doc,
        "Patterns that came up repeatedly across phases. Recording them "
        "here both because they will recur and because they affect how a "
        "reviewer should run the prototype locally.",
    )
    for item in OPERATIONS_LOG["items"]:
        add_h3(doc, item["name"])
        add_paragraph(doc, item["detail"])

    add_page_break(doc)

    add_h1(doc, CURRENT_STATE["title"])
    add_h2(doc, "Live and real in the prototype today")
    add_bullets(doc, CURRENT_STATE["live"])
    add_h2(doc, "Mocked, with documented production shape")
    add_bullets(doc, CURRENT_STATE["mocked"])
    add_h2(doc, "Deferred to a dedicated session, with a runbook checked in")
    add_bullets(doc, CURRENT_STATE["deferred_with_runbook"])

    add_h2(doc, "Closing note")
    add_paragraph(
        doc,
        "The prototype is intentionally honest about every mock. Every "
        "mocked API response carries a meta._note explaining what it is "
        "and what replaces it in production; every investor page includes "
        "a prototype-vs-production table; every Agent Fabric trace span "
        "shows _source: 'real' or 'mock' so the demo never lies about "
        "provenance even when one path is live. A reviewer can run the "
        "full intake -> A2A -> Care Router -> Data 360 grounding flow "
        "locally with no credentials and see the four-span trace in the "
        "live Agent Fabric console; the same reviewer with the three "
        "Salesforce env vars set sees the federated-query span flip to "
        "real Health Cloud data without any code change.",
    )
    add_paragraph(
        doc,
        "The three deferrals the original Phase 14 / 15 / 16 sessions "
        "ended with — Salesforce Data Cloud activation, Pause-owned "
        "Agentforce Embedded Messaging deployment, and the first live "
        "MuleSoft Anypoint Experience API — have ALL SHIPPED. Phase 18 "
        "stood up the Pause-Health-owned Agentforce deployment with the "
        "real Pause_Health_Intake_Agent. Phase 22 deployed "
        "pause-mulesoft-health-v1 to CloudHub 2.0 with live-or-mock "
        "graceful degradation; Phase 23 layered Flex Gateway runtime "
        "enforcement + JWT auth + rate limiting + an OAS 3.0 spec on "
        "Exchange. Phase 24 activated Data Cloud Phase 2 with three "
        "Calculated Insights live on trailsignup org. The current "
        "deferrals are smaller and more specific: Phase 1c Mule "
        "implementations of the 8 Phase-3 spec assets, the Phase-1 "
        "iteration-8 persistent VM that hosts Flex Gateway off the "
        "laptop, and the operator-side env-var procurement that flips "
        "the four Headless 360 audit gaps from prototype to shipped. "
        "Each is one focused session away.",
    )


def main() -> int:
    doc = Document()
    # set wider page (still letter), comfortable margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    render(doc)
    doc.save(OUTPUT_PATH)

    size_kb = OUTPUT_PATH.stat().st_size / 1024
    print(f"wrote {OUTPUT_PATH} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
